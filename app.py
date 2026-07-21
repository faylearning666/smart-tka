import sqlite3
try:
    import psycopg2
except ImportError:
    psycopg2 = None
import pandas as pd
import streamlit as st
from openai import OpenAI
import os
import PyPDF2
import docx
import json
import re
import html
import uuid

from io import BytesIO
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from datetime import date

import base64
import warnings

warnings.filterwarnings(
    "ignore",
    message="pandas only supports SQLAlchemy connectable.*",
    category=UserWarning
)

DB_NAME = "tka_mvp.db"

MAPEL_OPTIONS = [
    "Matematika",
    "Bahasa Indonesia",
    "Bahasa Inggris"
]

LEVEL_OPTIONS = ["Mudah", "Sedang", "Sulit"]
DB_SCHEMA_VERSION = "2026-07-21-v1"

st.set_page_config(
    page_title="TKA Digital MVP + LLM",
    page_icon="📘",
    layout="wide"
)


# =========================
# DATABASE
# =========================

def get_database_url():
    """
    Kalau DATABASE_URL diisi di Streamlit Secrets, aplikasi memakai Supabase/PostgreSQL.
    Kalau kosong, aplikasi tetap bisa jalan lokal memakai SQLite tka_mvp.db.
    """
    try:
        database_url = st.secrets.get("DATABASE_URL", None)
    except Exception:
        database_url = None

    if not database_url:
        database_url = os.environ.get("DATABASE_URL", None)

    return database_url


def is_postgres_mode():
    return bool(get_database_url())


def convert_sql_for_postgres(sql):
    """
    Adapter ringan supaya query lama SQLite lebih mudah jalan di PostgreSQL.
    """
    sql = sql.replace("INTEGER PRIMARY KEY AUTOINCREMENT", "SERIAL PRIMARY KEY")
    sql = sql.replace("AUTOINCREMENT", "")

    # SQLite pakai ?, PostgreSQL/psycopg2 pakai %s
    sql = sql.replace("?", "%s")

    return sql


def insert_needs_returning_id(sql):
    sql_clean = sql.strip().lower()
    return (
        sql_clean.startswith("insert into")
        and " returning " not in sql_clean
    )


class PostgresCursorWrapper:
    def __init__(self, cursor):
        self.cursor = cursor
        self.lastrowid = None

    def execute(self, sql, params=None):
        sql_pg = convert_sql_for_postgres(sql)

        needs_returning = insert_needs_returning_id(sql_pg)

        if needs_returning:
            sql_pg = sql_pg.rstrip().rstrip(";") + " RETURNING id"

        if params is None:
            self.cursor.execute(sql_pg)
        else:
            self.cursor.execute(sql_pg, params)

        if needs_returning:
            try:
                row = self.cursor.fetchone()
                self.lastrowid = row[0] if row else None
            except Exception:
                self.lastrowid = None

        return self

    def executemany(self, sql, seq_of_params):
        sql_pg = convert_sql_for_postgres(sql)
        self.cursor.executemany(sql_pg, seq_of_params)
        return self

    def fetchone(self):
        return self.cursor.fetchone()

    def fetchall(self):
        return self.cursor.fetchall()

    def close(self):
        return self.cursor.close()

    @property
    def description(self):
        return self.cursor.description

    @property
    def rowcount(self):
        return self.cursor.rowcount


class PostgresConnectionWrapper:
    def __init__(self, conn):
        self.conn = conn

    def cursor(self):
        return PostgresCursorWrapper(self.conn.cursor())

    def commit(self):
        return self.conn.commit()

    def rollback(self):
        return self.conn.rollback()

    def close(self):
        return self.conn.close()

    def __getattr__(self, name):
        return getattr(self.conn, name)


def connect_db():
    database_url = get_database_url()

    if database_url:
        if psycopg2 is None:
            raise ImportError(
                "psycopg2-binary belum terinstall. Tambahkan psycopg2-binary ke requirements.txt."
            )

        if "sslmode=" in database_url:
            conn = psycopg2.connect(
                database_url,
                connect_timeout=10
            )
        else:
            conn = psycopg2.connect(
                database_url,
                sslmode="require",
                connect_timeout=10
            )

        return PostgresConnectionWrapper(conn)

    return sqlite3.connect(DB_NAME, check_same_thread=False)


def tambah_kolom_jika_belum_ada(cur, nama_tabel, nama_kolom, tipe_data):
    if isinstance(cur, PostgresCursorWrapper):
        # Cek tabel di PostgreSQL
        cur.execute("""
        SELECT COUNT(*)
        FROM information_schema.tables
        WHERE table_schema = 'public'
        AND table_name = %s
        """, (nama_tabel,))

        if cur.fetchone()[0] == 0:
            return

        # Cek kolom di PostgreSQL
        cur.execute("""
        SELECT column_name
        FROM information_schema.columns
        WHERE table_schema = 'public'
        AND table_name = %s
        """, (nama_tabel,))

        kolom = [row[0] for row in cur.fetchall()]

        if nama_kolom not in kolom:
            cur.execute(f"ALTER TABLE {nama_tabel} ADD COLUMN {nama_kolom} {tipe_data}")

    else:
        # Cek tabel di SQLite
        cur.execute("""
        SELECT name
        FROM sqlite_master
        WHERE type='table' AND name=?
        """, (nama_tabel,))

        if cur.fetchone() is None:
            return

        # Cek kolom di SQLite
        cur.execute(f"PRAGMA table_info({nama_tabel})")
        kolom = [row[1] for row in cur.fetchall()]

        if nama_kolom not in kolom:
            cur.execute(f"ALTER TABLE {nama_tabel} ADD COLUMN {nama_kolom} {tipe_data}")


def init_db():
    conn = connect_db()
    cur = conn.cursor()

    # Khusus PostgreSQL/Supabase:
    # cegah init_db nyangkut terlalu lama saat Streamlit rerun/redeploy
    if isinstance(cur, PostgresCursorWrapper):
        cur.execute("SET lock_timeout TO '5s'")
        cur.execute("SET statement_timeout TO '20s'")

        cur.execute("SELECT pg_try_advisory_xact_lock(%s)", (2026072101,))
        got_lock = cur.fetchone()[0]

        if not got_lock:
            conn.close()
            return

    # =========================
    # 1. TABEL USER & ROLE
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE,
        password TEXT,
        role TEXT,
        nama TEXT,
        status_akun TEXT DEFAULT 'aktif'
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS admin (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        nama TEXT,
        email TEXT,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS siswa (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        nama TEXT,
        kelas TEXT,
        sekolah TEXT,
        email TEXT,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS ortu (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        siswa_id INTEGER,
        nama TEXT,
        hubungan TEXT,
        email TEXT,
        no_hp TEXT,
        FOREIGN KEY (user_id) REFERENCES users(id),
        FOREIGN KEY (siswa_id) REFERENCES siswa(id)
    )
    """)

    # =========================
    # 2. TABEL TRY OUT
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS hasil_tryout (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        siswa_id INTEGER,
        mapel TEXT,
        nilai REAL,
        benar INTEGER,
        total INTEGER,
        rekomendasi TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (siswa_id) REFERENCES siswa(id)
    )
    """)


    # Bank soal dibuat lebih awal karena hasil_tryout_detail punya foreign key ke bank_soal.
    # Ini penting untuk PostgreSQL/Supabase.
    cur.execute("""
    CREATE TABLE IF NOT EXISTS bank_soal (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mapel TEXT,
        topik TEXT,
        level TEXT,
        pertanyaan TEXT,
        opsi_a TEXT,
        opsi_b TEXT,
        opsi_c TEXT,
        opsi_d TEXT,
        jawaban TEXT,
        pembahasan TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS hasil_tryout_detail (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        hasil_tryout_id INTEGER,
        siswa_id INTEGER,
        soal_id INTEGER,
        mapel TEXT,
        topik TEXT,
        level TEXT,
        jawaban_siswa TEXT,
        jawaban_benar TEXT,
        is_benar INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (hasil_tryout_id) REFERENCES hasil_tryout(id),
        FOREIGN KEY (siswa_id) REFERENCES siswa(id),
        FOREIGN KEY (soal_id) REFERENCES bank_soal(id)
    )
    """)

    # =========================
    # 3. TABEL BANK SOAL
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS bank_soal (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mapel TEXT,
        topik TEXT,
        level TEXT,
        pertanyaan TEXT,
        opsi_a TEXT,
        opsi_b TEXT,
        opsi_c TEXT,
        opsi_d TEXT,
        jawaban TEXT,
        pembahasan TEXT
    )
    """)

    # =========================
    # 4. TABEL MATERI & CHUNK
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS materi (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        judul TEXT,
        filename TEXT,
        isi_teks TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS materi_chunk (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        materi_id INTEGER,
        judul TEXT,
        isi_chunk TEXT,
        urutan INTEGER,
        FOREIGN KEY (materi_id) REFERENCES materi(id)
    )
    """)

    # =========================
    # 5. TABEL REGISTRASI ORTU & SISWA
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS pendaftaran_ortu_siswa (
        id INTEGER PRIMARY KEY AUTOINCREMENT,

        nama_siswa TEXT,
        kelas TEXT,
        sekolah TEXT,
        email_siswa TEXT,
        username_siswa TEXT,
        password_siswa TEXT,

        nama_ortu TEXT,
        hubungan TEXT,
        email_ortu TEXT,
        no_hp TEXT,
        username_ortu TEXT,
        password_ortu TEXT,

        status TEXT DEFAULT 'pending',
        catatan_admin TEXT,

        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP
    )
    """)

    # =========================
    # 6. TABEL JADWAL BELAJAR
    # =========================
    cur.execute("""
    CREATE TABLE IF NOT EXISTS jadwal_belajar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        siswa_id INTEGER,
        mapel TEXT,
        tanggal_mulai TEXT,
        jumlah_hari INTEGER,
        target_nilai INTEGER,
        waktu_per_hari TEXT,
        jam_belajar TEXT,
        reminder_aktif TEXT,
        catatan_ai TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (siswa_id) REFERENCES siswa(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS jadwal_belajar_detail (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        jadwal_id INTEGER,
        hari_ke INTEGER,
        tanggal TEXT,
        topik TEXT,
        level TEXT,
        topik_spesifik TEXT,
        konteks_kelemahan TEXT,
        fokus_belajar TEXT,
        materi_harian TEXT,
        latihan_disarankan TEXT,
        aktivitas TEXT,
        status TEXT DEFAULT 'belum',
        FOREIGN KEY (jadwal_id) REFERENCES jadwal_belajar(id)
    )
    """)

    # TABEL VISITOR WEB
    cur.execute("""
    CREATE TABLE IF NOT EXISTS visitor_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        session_id TEXT,
        halaman TEXT,
        role TEXT,
        username TEXT,
        tanggal TEXT,
        is_login INTEGER DEFAULT 0,
        login_at TIMESTAMP,
        waktu TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)

    # =========================
    # TABEL FREE TRIAL LOG
    # =========================
    
    if isinstance(cur, PostgresCursorWrapper):
        # Cek dulu apakah tabel sudah ada di Supabase/PostgreSQL
        cur.execute("SELECT to_regclass(%s)", ("public.free_trial_log",))
        table_exists = cur.fetchone()[0] is not None
    
        if not table_exists:
            cur.execute("""
            CREATE TABLE free_trial_log (
                id BIGSERIAL PRIMARY KEY,
                session_id TEXT,
                nama TEXT,
                email TEXT,
                no_hp TEXT,
                mapel TEXT,
                nilai REAL,
                benar INTEGER,
                total INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """)
    
    else:
        # Versi SQLite lokal
        cur.execute("""
        CREATE TABLE IF NOT EXISTS free_trial_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT,
            nama TEXT,
            email TEXT,
            no_hp TEXT,
            mapel TEXT,
            nilai REAL,
            benar INTEGER,
            total INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)

    # =========================
    # 7. MIGRASI AMAN UNTUK DATABASE LAMA
    # =========================

    # Users
    tambah_kolom_jika_belum_ada(cur, "users", "status_akun", "TEXT")

    # Bank soal
    tambah_kolom_jika_belum_ada(cur, "bank_soal", "topik", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "bank_soal", "level", "TEXT")

    # Hasil tryout
    tambah_kolom_jika_belum_ada(cur, "hasil_tryout", "mapel", "TEXT")

    # Jadwal belajar
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar", "jam_belajar", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar", "reminder_aktif", "TEXT")

    # Jadwal belajar detail
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar_detail", "fokus_belajar", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar_detail", "materi_harian", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar_detail", "latihan_disarankan", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar_detail", "topik_spesifik", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "jadwal_belajar_detail", "konteks_kelemahan", "TEXT")

    tambah_kolom_jika_belum_ada(cur, "visitor_log", "is_", "INTEGER DEFAULT 0")
    tambah_kolom_jika_belum_ada(cur, "visitor_log", "_at", "TIMESTAMP")

    # tambah untuk pembayaran
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "jenis_pendaftaran", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "nominal_bayar", "INTEGER")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "status_pembayaran", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "bukti_bayar_filename", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "bukti_bayar_mime", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "bukti_bayar_base64", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "sumber_trial_email", "TEXT")
    tambah_kolom_jika_belum_ada(cur, "pendaftaran_ortu_siswa", "sumber_trial_no_hp", "TEXT")

    # Pastikan akun lama aktif
    cur.execute("""
    UPDATE users 
    SET status_akun = 'aktif' 
    WHERE status_akun IS NULL OR status_akun = ''
    """)

    # =========================
    # 8. SEED USER DEMO
    # =========================
    cur.execute("SELECT COUNT(*) FROM users")

    if cur.fetchone()[0] == 0:
        cur.executemany("""
        INSERT INTO users (username, password, role, nama, status_akun)
        VALUES (?, ?, ?, ?, ?)
        """, [
            ("admin", "admin123", "admin", "Admin", "aktif"),
            ("siswa", "siswa123", "siswa", "Siswa Demo", "aktif"),
            ("ortu", "ortu123", "ortu", "Orang Tua Demo", "aktif"),
        ])

    # =========================
    # 9. SEED PROFIL DEMO SISWA & ORTU
    # =========================

    # Ambil user siswa demo
    cur.execute("SELECT id FROM users WHERE username='siswa'")
    row_siswa_user = cur.fetchone()

    if row_siswa_user:
        user_siswa_id = row_siswa_user[0]

        cur.execute("SELECT COUNT(*) FROM siswa WHERE user_id=?", (user_siswa_id,))
        if cur.fetchone()[0] == 0:
            cur.execute("""
            INSERT INTO siswa (user_id, nama, kelas, sekolah, email)
            VALUES (?, ?, ?, ?, ?)
            """, (
                user_siswa_id,
                "Siswa Demo",
                "3",
                "Sekolah Demo",
                "siswa.demo@email.com"
            ))

    # Ambil data siswa demo setelah dibuat
    cur.execute("""
    SELECT siswa.id
    FROM siswa
    JOIN users ON siswa.user_id = users.id
    WHERE users.username='siswa'
    """)
    row_siswa = cur.fetchone()

    cur.execute("SELECT id FROM users WHERE username='ortu'")
    row_ortu_user = cur.fetchone()

    if row_siswa and row_ortu_user:
        siswa_demo_id = row_siswa[0]
        user_ortu_id = row_ortu_user[0]

        cur.execute("SELECT COUNT(*) FROM ortu WHERE user_id=?", (user_ortu_id,))
        if cur.fetchone()[0] == 0:
            cur.execute("""
            INSERT INTO ortu (user_id, siswa_id, nama, hubungan, email, no_hp)
            VALUES (?, ?, ?, ?, ?, ?)
            """, (
                user_ortu_id,
                siswa_demo_id,
                "Orang Tua Demo",
                "Ibu",
                "ortu.demo@email.com",
                "081234567890"
            ))

    # =========================
    # 10. SEED BANK SOAL DEMO
    # =========================
    cur.execute("SELECT COUNT(*) FROM bank_soal")

    if cur.fetchone()[0] == 0:
        cur.executemany("""
        INSERT INTO bank_soal
        (mapel, topik, level, pertanyaan, opsi_a, opsi_b, opsi_c, opsi_d, jawaban, pembahasan)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, [
            (
                "Matematika",
                "Perkalian Dasar",
                "Mudah",
                "Hasil dari 12 x 8 adalah...",
                "80",
                "88",
                "96",
                "108",
                "C",
                "12 x 8 = 96."
            ),
            (
                "Bahasa Indonesia",
                "Sinonim",
                "Mudah",
                "Sinonim dari kata efektif adalah...",
                "Tepat guna",
                "Lambat",
                "Rumit",
                "Sia-sia",
                "A",
                "Efektif berarti berhasil guna atau tepat guna."
            ),
            (
                "Bahasa Inggris",
                "Vocabulary",
                "Mudah",
                "What is the meaning of 'book' in Indonesian?",
                "Meja",
                "Buku",
                "Kursi",
                "Pensil",
                "B",
                "'Book' berarti 'buku' dalam bahasa Indonesia."
            )
        ])

    # =========================
    # 11. MIGRASI ISI DATA LAMA
    # =========================
    cur.execute("UPDATE bank_soal SET mapel='Matematika' WHERE mapel='Numerasi'")
    cur.execute("UPDATE bank_soal SET mapel='Bahasa Indonesia' WHERE mapel='Literasi'")
    cur.execute("UPDATE bank_soal SET mapel='Bahasa Indonesia' WHERE mapel='Penalaran Umum'")

    conn.commit()
    conn.close()

@st.cache_resource(show_spinner=False)
def jalankan_init_db_sekali(schema_version):
    init_db()
    return True
    
#================================
# cek trial gratis
#================================
HARGA_REGISTRASI = 100000

def normalisasi_teks(teks):
    if teks is None:
        return ""
    return str(teks).strip().lower()


def sudah_pernah_tryout_gratis(email, no_hp):
    email = normalisasi_teks(email)
    no_hp = normalisasi_teks(no_hp)

    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT COUNT(*)
    FROM free_trial_log
    WHERE LOWER(email) = ? OR no_hp = ?
    """, (email, no_hp))

    jumlah = cur.fetchone()[0]
    conn.close()

    return jumlah > 0


def catat_tryout_gratis(nama, email, no_hp, mapel, nilai, benar, total):
    if "visitor_session_id" not in st.session_state:
        st.session_state["visitor_session_id"] = str(uuid.uuid4())

    session_id = st.session_state["visitor_session_id"]

    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    INSERT INTO free_trial_log
    (session_id, nama, email, no_hp, mapel, nilai, benar, total)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        session_id,
        nama,
        normalisasi_teks(email),
        no_hp,
        mapel,
        nilai,
        benar,
        total
    ))

    conn.commit()
    conn.close()

# =========================
# helper buat redirect halaman
# =========================
def redirect_ke_register_bayar():
    st.session_state["free_soal_ids"] = []
    st.session_state["free_trial_identitas_ok"] = False
    st.session_state["halaman_awal"] = "register_bayar"

    # Bersihkan jawaban radio tryout gratis supaya tidak nyangkut
    keys_to_delete = [
        key for key in st.session_state.keys()
        if str(key).startswith("free_soal_")
    ]

    for key in keys_to_delete:
        del st.session_state[key]

# =========================
# FUNGSI NAIK TURUN LEVEL UNTUK ADAPTIVE LEARNING
# =========================
def naik_level(level):
    if level == "Mudah":
        return "Sedang"
    elif level == "Sedang":
        return "Sulit"
    else:
        return "Sulit"


def turun_level(level):
    if level == "Sulit":
        return "Sedang"
    elif level == "Sedang":
        return "Mudah"
    else:
        return "Mudah"


def tentukan_level_berikutnya(akurasi, level_terakhir, total_dikerjakan):
    if not level_terakhir:
        level_terakhir = "Mudah"

    # Kalau datanya masih sedikit, jangan terlalu agresif naik/turun
    if total_dikerjakan < 3:
        return level_terakhir

    if akurasi >= 80:
        return naik_level(level_terakhir)

    elif akurasi < 50:
        return turun_level(level_terakhir)

    else:
        return level_terakhir


# =========================
# AMBIL SOAL ADAPTIVE LEARNING
# =========================
def ambil_soal_adaptif(siswa_id, mapel, jumlah_soal=10):
    conn = connect_db()

    df_soal_all = pd.read_sql_query("""
    SELECT *
    FROM bank_soal
    WHERE mapel = ?
    """, conn, params=(mapel,))

    if df_soal_all.empty:
        conn.close()
        return pd.DataFrame(), pd.DataFrame()

    df_soal_all["topik"] = df_soal_all["topik"].fillna("")
    df_soal_all["level"] = df_soal_all["level"].fillna("Mudah")

    df_performa = pd.read_sql_query("""
    SELECT 
        topik,
        COUNT(*) AS total_dikerjakan,
        SUM(CASE WHEN is_benar = 1 THEN 1 ELSE 0 END) AS total_benar,
        ROUND(AVG(is_benar) * 100, 2) AS akurasi
    FROM hasil_tryout_detail
    WHERE siswa_id = ? AND mapel = ?
    GROUP BY topik
    """, conn, params=(siswa_id, mapel))

    df_level_terakhir = pd.read_sql_query("""
    SELECT topik, level, created_at
    FROM hasil_tryout_detail
    WHERE siswa_id = ? AND mapel = ?
    ORDER BY created_at DESC
    """, conn, params=(siswa_id, mapel))

    df_recent = pd.read_sql_query("""
    SELECT soal_id
    FROM hasil_tryout_detail
    WHERE siswa_id = ? AND mapel = ?
    ORDER BY created_at DESC
    LIMIT 50
    """, conn, params=(siswa_id, mapel))

    conn.close()

    recent_soal_ids = df_recent["soal_id"].tolist() if not df_recent.empty else []

    topik_list = df_soal_all["topik"].dropna().unique().tolist()

    if not topik_list:
        topik_list = [""]

    last_level_map = {}

    for _, row in df_level_terakhir.iterrows():
        if row["topik"] not in last_level_map:
            last_level_map[row["topik"]] = row["level"]

    info_adaptif = []
    soal_terpilih = []

    for topik in topik_list:
        perf = df_performa[df_performa["topik"] == topik]

        if perf.empty:
            total_dikerjakan = 0
            akurasi = 0
            level_terakhir = "Mudah"
            level_target = "Mudah"
        else:
            total_dikerjakan = int(perf.iloc[0]["total_dikerjakan"])
            akurasi = float(perf.iloc[0]["akurasi"])
            level_terakhir = last_level_map.get(topik, "Mudah")
            level_target = tentukan_level_berikutnya(
                akurasi,
                level_terakhir,
                total_dikerjakan
            )

        kandidat = df_soal_all[
            (df_soal_all["topik"] == topik) &
            (df_soal_all["level"] == level_target) &
            (~df_soal_all["id"].isin(recent_soal_ids))
            ]

        # fallback 1: topik sama, level bebas, belum pernah baru-baru ini
        if kandidat.empty:
            kandidat = df_soal_all[
                (df_soal_all["topik"] == topik) &
                (~df_soal_all["id"].isin(recent_soal_ids))
                ]

        # fallback 2: topik sama, soal apa saja
        if kandidat.empty:
            kandidat = df_soal_all[
                df_soal_all["topik"] == topik
                ]

        if not kandidat.empty:
            soal_terpilih.append(kandidat.sample(1).iloc[0])

        info_adaptif.append({
            "topik": topik,
            "total_dikerjakan": total_dikerjakan,
            "akurasi": akurasi,
            "level_terakhir": level_terakhir,
            "level_soal_berikutnya": level_target
        })

    df_terpilih = pd.DataFrame(soal_terpilih)

    if len(df_terpilih) < jumlah_soal:
        id_terpilih = df_terpilih["id"].tolist() if not df_terpilih.empty else []

        tambahan = df_soal_all[
            ~df_soal_all["id"].isin(id_terpilih + recent_soal_ids)
        ]

        if tambahan.empty:
            tambahan = df_soal_all[
                ~df_soal_all["id"].isin(id_terpilih)
            ]

        sisa = jumlah_soal - len(df_terpilih)

        if not tambahan.empty:
            tambahan = tambahan.sample(
                min(sisa, len(tambahan))
            )

            df_terpilih = pd.concat(
                [df_terpilih, tambahan],
                ignore_index=True
            )

    df_terpilih = df_terpilih.head(jumlah_soal)

    df_info = pd.DataFrame(info_adaptif)

    return df_terpilih, df_info


# =========================
# LLM
# =========================

def get_llm_client():
    api_key = st.secrets.get("OPENROUTER_API_KEY", None)

    if not api_key:
        return None

    return OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key
    )


def ask_llm(prompt):
    client = get_llm_client()

    if client is None:
        return "LLM belum aktif. Tambahkan OPENROUTER_API_KEY di file .streamlit/secrets.toml."

    response = client.chat.completions.create(
        model="openai/gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "Kamu adalah tutor TKA yang menjelaskan materi dengan bahasa Indonesia yang sederhana, runtut, dan ramah."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.4
    )

    return response.choices[0].message.content


# =========================
# AUTH
# =========================

def login(username, password):
    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT username, role, nama, status_akun
    FROM users
    WHERE username = ? AND password = ?
    """, (username, password))

    user = cur.fetchone()
    conn.close()

    if user is None:
        return None

    if user[3] != "aktif":
        return "nonaktif"

    return user


# =========================
# ambil ID siswa dan ortu
# =========================
def get_siswa_id_by_username(username):
    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT siswa.id
    FROM siswa
    JOIN users ON siswa.user_id = users.id
    WHERE users.username = ?
    """, (username,))

    row = cur.fetchone()
    conn.close()

    return row[0] if row else None


def get_siswa_id_by_ortu_username(username):
    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT ortu.siswa_id
    FROM ortu
    JOIN users ON ortu.user_id = users.id
    WHERE users.username = ?
    """, (username,))

    row = cur.fetchone()
    conn.close()

    return row[0] if row else None


# =====================================
# hitung berapa kali TKA
# =====================================
def ambil_rekap_simulasi_tka(siswa_id=None):
    conn = connect_db()

    query = """
    SELECT 
        siswa.id AS siswa_id,
        siswa.nama AS nama_siswa,
        siswa.kelas,
        siswa.sekolah,
        users.status_akun,
        COUNT(hasil_tryout.id) AS total_simulasi,
        AVG(hasil_tryout.nilai) AS rata_nilai,
        MAX(hasil_tryout.nilai) AS nilai_tertinggi,
        MAX(hasil_tryout.created_at) AS terakhir_simulasi
    FROM siswa
    JOIN users ON siswa.user_id = users.id
    LEFT JOIN hasil_tryout ON hasil_tryout.siswa_id = siswa.id
    WHERE users.status_akun = 'aktif'
    """

    params = []

    if siswa_id is not None:
        query += " AND siswa.id = ?"
        params.append(siswa_id)

    query += """
    GROUP BY 
        siswa.id,
        siswa.nama,
        siswa.kelas,
        siswa.sekolah,
        users.status_akun
    ORDER BY total_simulasi DESC, siswa.nama ASC
    """

    df = pd.read_sql_query(query, conn, params=params)
    conn.close()

    if not df.empty:
        df["total_simulasi"] = df["total_simulasi"].fillna(0).astype(int)
        df["rata_nilai"] = pd.to_numeric(df["rata_nilai"], errors="coerce").fillna(0).round(2)
        df["nilai_tertinggi"] = pd.to_numeric(df["nilai_tertinggi"], errors="coerce").fillna(0).round(2)
        df["terakhir_simulasi"] = df["terakhir_simulasi"].fillna("-")

    return df


def ambil_rekap_simulasi_per_mapel(siswa_id):
    conn = connect_db()

    df = pd.read_sql_query("""
    SELECT 
        mapel,
        COUNT(*) AS jumlah_simulasi,
        AVG(nilai) AS rata_nilai,
        MAX(nilai) AS nilai_tertinggi,
        MAX(created_at) AS terakhir_simulasi
    FROM hasil_tryout
    WHERE siswa_id = ?
    GROUP BY mapel
    ORDER BY jumlah_simulasi DESC, mapel ASC
    """, conn, params=(siswa_id,))

    conn.close()

    if not df.empty:
        df["jumlah_simulasi"] = df["jumlah_simulasi"].fillna(0).astype(int)
        df["rata_nilai"] = pd.to_numeric(df["rata_nilai"], errors="coerce").fillna(0).round(2)
        df["nilai_tertinggi"] = pd.to_numeric(df["nilai_tertinggi"], errors="coerce").fillna(0).round(2)
        df["terakhir_simulasi"] = df["terakhir_simulasi"].fillna("-")

    return df


def tampilkan_rekap_penggunaan_simulasi_admin():
    st.subheader("📌 Rekap Penggunaan Simulasi TKA")

    df_rekap = ambil_rekap_simulasi_tka()

    if df_rekap.empty:
        st.info("Belum ada data siswa aktif.")
        return

    total_siswa = len(df_rekap)
    siswa_sudah_simulasi = int((df_rekap["total_simulasi"] > 0).sum())
    total_simulasi = int(df_rekap["total_simulasi"].sum())
    rata_pemakaian = round(total_simulasi / total_siswa, 2) if total_siswa > 0 else 0

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Siswa Aktif", total_siswa)

    with col2:
        st.metric("Sudah Simulasi", siswa_sudah_simulasi)

    with col3:
        st.metric("Total Simulasi", total_simulasi)

    with col4:
        st.metric("Rata-rata / Siswa", rata_pemakaian)

    df_tampil = df_rekap.rename(columns={
        "siswa_id": "ID Siswa",
        "nama_siswa": "Nama Siswa",
        "kelas": "Kelas",
        "sekolah": "Sekolah",
        "total_simulasi": "Total Simulasi",
        "rata_nilai": "Rata-rata Nilai",
        "nilai_tertinggi": "Nilai Tertinggi",
        "terakhir_simulasi": "Terakhir Simulasi"
    })

    kolom_tampil = [
        "ID Siswa",
        "Nama Siswa",
        "Kelas",
        "Sekolah",
        "Total Simulasi",
        "Rata-rata Nilai",
        "Nilai Tertinggi",
        "Terakhir Simulasi"
    ]

    st.dataframe(df_tampil[kolom_tampil], use_container_width=True)


def tampilkan_penggunaan_simulasi_siswa(siswa_id, judul="📌 Penggunaan Simulasi TKA"):
    st.subheader(judul)

    df_rekap = ambil_rekap_simulasi_tka(siswa_id)

    if df_rekap.empty:
        st.info("Data siswa tidak ditemukan.")
        return

    data = df_rekap.iloc[0]

    total_simulasi = int(data["total_simulasi"])
    rata_nilai = float(data["rata_nilai"])
    nilai_tertinggi = float(data["nilai_tertinggi"])
    terakhir_simulasi = data["terakhir_simulasi"]

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("Total Simulasi TKA", total_simulasi)

    with col2:
        st.metric("Rata-rata Nilai", rata_nilai)

    with col3:
        st.metric("Nilai Tertinggi", nilai_tertinggi)

    st.caption(f"Terakhir simulasi: {terakhir_simulasi}")

    df_mapel = ambil_rekap_simulasi_per_mapel(siswa_id)

    if not df_mapel.empty:
        st.write("**Rincian per Mapel**")

        df_mapel_tampil = df_mapel.rename(columns={
            "mapel": "Mapel",
            "jumlah_simulasi": "Jumlah Simulasi",
            "rata_nilai": "Rata-rata Nilai",
            "nilai_tertinggi": "Nilai Tertinggi",
            "terakhir_simulasi": "Terakhir Simulasi"
        })

        st.dataframe(df_mapel_tampil, use_container_width=True)
    else:
        st.info("Belum pernah melakukan simulasi TKA.")
        
# =================
# FUNGSI BUAT NGECEK USERNAME
# =================
def username_sudah_dipakai(username):
    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT COUNT(*)
    FROM users
    WHERE username = ?
    """, (username,))

    jumlah_users = cur.fetchone()[0]

    cur.execute("""
    SELECT COUNT(*)
    FROM pendaftaran_ortu_siswa
    WHERE (username_siswa = ? OR username_ortu = ?)
    AND status = 'pending'
    """, (username, username))

    jumlah_pending = cur.fetchone()[0]

    conn.close()

    return jumlah_users > 0 or jumlah_pending > 0


# =================
# HALAMAN REGISTER ORANG TUA
# =================
def page_register_ortu():
    st.title("📝 Registrasi Orang Tua & Siswa")

    st.info("Data registrasi akan diverifikasi terlebih dahulu oleh admin sebelum akun dapat digunakan.")

    with st.form("form_register_ortu_siswa"):
        st.subheader("Data Siswa")

        nama_siswa = st.text_input("Nama Siswa")
        kelas = st.text_input("Kelas")
        sekolah = st.text_input("Sekolah")
        email_siswa = st.text_input("Email Siswa")

        username_siswa = st.text_input("Username Login Siswa")
        password_siswa = st.text_input("Password Login Siswa", type="password")

        st.divider()

        st.subheader("Data Orang Tua / Wali")

        nama_ortu = st.text_input("Nama Orang Tua / Wali")
        hubungan = st.selectbox("Hubungan", ["Ayah", "Ibu", "Wali"])
        email_ortu = st.text_input("Email Orang Tua")
        no_hp = st.text_input("No HP")

        username_ortu = st.text_input("Username Login Orang Tua")
        password_ortu = st.text_input("Password Login Orang Tua", type="password")

        submit = st.form_submit_button("Kirim Registrasi")

        if submit:
            if not nama_siswa or not nama_ortu or not username_siswa or not username_ortu or not password_siswa or not password_ortu:
                st.warning("Nama siswa, nama orang tua, username, dan password wajib diisi.")
                return

            if username_siswa == username_ortu:
                st.warning("Username siswa dan username orang tua tidak boleh sama.")
                return

            if username_sudah_dipakai(username_siswa):
                st.error("Username siswa sudah digunakan atau sedang menunggu verifikasi.")
                return

            if username_sudah_dipakai(username_ortu):
                st.error("Username orang tua sudah digunakan atau sedang menunggu verifikasi.")
                return

            conn = connect_db()
            cur = conn.cursor()

            try:
                cur.execute("""
                INSERT INTO pendaftaran_ortu_siswa (
                    nama_siswa, kelas, sekolah, email_siswa, username_siswa, password_siswa,
                    nama_ortu, hubungan, email_ortu, no_hp, username_ortu, password_ortu,
                    status
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending')
                """, (
                    nama_siswa, kelas, sekolah, email_siswa, username_siswa, password_siswa,
                    nama_ortu, hubungan, email_ortu, no_hp, username_ortu, password_ortu
                ))

                conn.commit()
                st.success("Registrasi berhasil dikirim. Silakan menunggu verifikasi admin.")

            except Exception as e:
                conn.rollback()
                st.error("Registrasi gagal.")
                st.code(str(e))

            conn.close()


# ===================
# halaman registrasi bayar
# ===================
def page_register_bayar():
    st.title("💳 Registrasi Berbayar Smart TKA")

    st.success("Aktifkan akun penuh dan akses semua fitur Smart TKA.")
    st.info(
        f"Biaya registrasi: **Rp {HARGA_REGISTRASI:,.0f}**".replace(",", ".")
        + "\n\nSilakan lakukan pembayaran via transfer ke No. Rek 202604606 BNI a.n. Diana Effendi, lalu upload bukti bayar."
    )

    st.warning(
        "Setelah registrasi dikirim, admin akan memverifikasi bukti bayar terlebih dahulu. "
        "Akun siswa dan orang tua akan aktif setelah disetujui admin."
    )

    with st.form("form_register_bayar"):
        st.subheader("Data Siswa")

        nama_siswa = st.text_input("Nama Siswa")
        kelas = st.text_input("Kelas")
        sekolah = st.text_input("Sekolah")
        email_siswa = st.text_input(
            "Email Siswa",
            value=st.session_state.get("trial_email", "")
        )

        username_siswa = st.text_input("Username Login Siswa")
        password_siswa = st.text_input("Password Login Siswa", type="password")

        st.divider()

        st.subheader("Data Orang Tua / Wali")

        nama_ortu = st.text_input("Nama Orang Tua / Wali")
        hubungan = st.selectbox("Hubungan", ["Ayah", "Ibu", "Wali"])
        email_ortu = st.text_input("Email Orang Tua")
        no_hp = st.text_input(
            "No HP",
            value=st.session_state.get("trial_no_hp", "")
        )

        username_ortu = st.text_input("Username Login Orang Tua")
        password_ortu = st.text_input("Password Login Orang Tua", type="password")

        st.divider()

        st.subheader("Upload Bukti Pembayaran")
        bukti_bayar = st.file_uploader(
            "Upload bukti bayar",
            type=["jpg", "jpeg", "png", "pdf"]
        )

        submit = st.form_submit_button("Kirim Registrasi Berbayar")

        if submit:
            if not nama_siswa or not nama_ortu or not username_siswa or not username_ortu:
                st.warning("Nama siswa, nama orang tua, username siswa, dan username orang tua wajib diisi.")
                return

            if not password_siswa or not password_ortu:
                st.warning("Password siswa dan password orang tua wajib diisi.")
                return

            if username_siswa == username_ortu:
                st.warning("Username siswa dan username orang tua tidak boleh sama.")
                return

            if username_sudah_dipakai(username_siswa):
                st.error("Username siswa sudah digunakan atau sedang menunggu verifikasi.")
                return

            if username_sudah_dipakai(username_ortu):
                st.error("Username orang tua sudah digunakan atau sedang menunggu verifikasi.")
                return

            if bukti_bayar is None:
                st.warning("Bukti bayar wajib diupload.")
                return

            bukti_bytes = bukti_bayar.read()
            bukti_base64 = base64.b64encode(bukti_bytes).decode("utf-8")

            conn = connect_db()
            cur = conn.cursor()

            try:
                cur.execute("""
                INSERT INTO pendaftaran_ortu_siswa (
                    nama_siswa, kelas, sekolah, email_siswa, username_siswa, password_siswa,
                    nama_ortu, hubungan, email_ortu, no_hp, username_ortu, password_ortu,
                    status, jenis_pendaftaran, nominal_bayar, status_pembayaran,
                    bukti_bayar_filename, bukti_bayar_mime, bukti_bayar_base64,
                    sumber_trial_email, sumber_trial_no_hp
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 
                        'pending', 'berbayar', ?, 'menunggu_verifikasi',
                        ?, ?, ?, ?, ?)
                """, (
                    nama_siswa, kelas, sekolah, email_siswa, username_siswa, password_siswa,
                    nama_ortu, hubungan, email_ortu, no_hp, username_ortu, password_ortu,
                    HARGA_REGISTRASI,
                    bukti_bayar.name,
                    bukti_bayar.type,
                    bukti_base64,
                    st.session_state.get("trial_email", ""),
                    st.session_state.get("trial_no_hp", "")
                ))

                conn.commit()

                st.success(
                    "Registrasi berbayar berhasil dikirim. "
                    "Silakan tunggu verifikasi admin."
                )

            except Exception as e:
                conn.rollback()
                st.error("Registrasi berbayar gagal.")
                st.code(str(e))

            conn.close()
            
# =================
# HALAMAN VERIFIKASI ADMIN NERIMA ORTU DAN SISWA
# =================
def page_verifikasi_register():
    st.title("✅ Verifikasi Registrasi")

    conn = connect_db()

    df = pd.read_sql_query("""
        SELECT *
        FROM pendaftaran_ortu_siswa
        ORDER BY created_at DESC
        """, conn)

    if df.empty:
        st.info("Belum ada data registrasi.")
        conn.close()
        return

    total_pending = len(df[df["status"] == "pending"])
    total_diterima = len(df[df["status"] == "diterima"])
    total_ditolak = len(df[df["status"] == "ditolak"])

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("Pending", total_pending)

    with col2:
        st.metric("Diterima", total_diterima)

    with col3:
        st.metric("Ditolak", total_ditolak)

    tab_pending, tab_diterima, tab_ditolak = st.tabs([
        "⏳ Pending",
        "✅ Diterima",
        "❌ Ditolak"
    ])

    with tab_pending:
        df_pending = df[df["status"] == "pending"]
        st.subheader("Registrasi Pending")
        st.dataframe(df_pending, use_container_width=True)

        if df_pending.empty:
            st.success("Tidak ada registrasi pending.")
        else:
            selected_id = st.selectbox(
                "Pilih ID Registrasi Pending",
                df_pending["id"].tolist()
            )

            data = df_pending[df_pending["id"] == selected_id].iloc[0]

            col_a, col_b = st.columns(2)

            with col_a:
                st.write("### Data Siswa")
                st.write(f"Nama: **{data['nama_siswa']}**")
                st.write(f"Kelas: **{data['kelas']}**")
                st.write(f"Sekolah: **{data['sekolah']}**")
                st.write(f"Email: **{data['email_siswa']}**")
                st.write(f"Username: **{data['username_siswa']}**")

            with col_b:
                st.write("### Data Orang Tua")
                st.write(f"Nama: **{data['nama_ortu']}**")
                st.write(f"Hubungan: **{data['hubungan']}**")
                st.write(f"Email: **{data['email_ortu']}**")
                st.write(f"No HP: **{data['no_hp']}**")
                st.write(f"Username: **{data['username_ortu']}**")

                st.write("### Info Pembayaran")
    
                jenis_pendaftaran = data["jenis_pendaftaran"] if "jenis_pendaftaran" in data and data["jenis_pendaftaran"] else "-"
                nominal_bayar = data["nominal_bayar"] if "nominal_bayar" in data and data["nominal_bayar"] else 0
                status_pembayaran = data["status_pembayaran"] if "status_pembayaran" in data and data["status_pembayaran"] else "-"
                
                st.write(f"Jenis pendaftaran: **{jenis_pendaftaran}**")
                st.write(f"Nominal bayar: **Rp {int(nominal_bayar):,}**".replace(",", "."))
                st.write(f"Status pembayaran: **{status_pembayaran}**")
                
                if "bukti_bayar_base64" in data and data["bukti_bayar_base64"]:
                    bukti_bytes = base64.b64decode(data["bukti_bayar_base64"])
                    filename = data["bukti_bayar_filename"] if data["bukti_bayar_filename"] else "bukti_bayar"
                    mime = data["bukti_bayar_mime"] if data["bukti_bayar_mime"] else "application/octet-stream"
                
                    st.download_button(
                        "Download Bukti Bayar",
                        data=bukti_bytes,
                        file_name=filename,
                        mime=mime
                    )
                else:
                    st.caption("Belum ada bukti bayar.")

            catatan_admin = st.text_area("Catatan Admin")

            col_acc, col_reject = st.columns(2)

            with col_acc:
                if st.button("Terima Registrasi", use_container_width=True):
                    proses_terima_registrasi(conn, data, selected_id, catatan_admin)

            with col_reject:
                if st.button("Tolak Registrasi", use_container_width=True):
                    proses_tolak_registrasi(conn, selected_id, catatan_admin)

    with tab_diterima:
        st.subheader("Registrasi Diterima")
        st.dataframe(df[df["status"] == "diterima"], use_container_width=True)

    with tab_ditolak:
        st.subheader("Registrasi Ditolak")
        st.dataframe(df[df["status"] == "ditolak"], use_container_width=True)

    conn.close()


# ================
# TERIMA REGISTRASI
# ================
def proses_terima_registrasi(conn, data, selected_id, catatan_admin):
    cur = conn.cursor()

    try:
        if username_sudah_dipakai(data["username_siswa"]) and data["status"] == "pending":
            cur.execute("""
            SELECT COUNT(*)
            FROM users
            WHERE username = ?
            """, (data["username_siswa"],))

            if cur.fetchone()[0] > 0:
                st.error("Username siswa sudah ada di tabel users.")
                conn.close()
                return

        if username_sudah_dipakai(data["username_ortu"]) and data["status"] == "pending":
            cur.execute("""
            SELECT COUNT(*)
            FROM users
            WHERE username = ?
            """, (data["username_ortu"],))

            if cur.fetchone()[0] > 0:
                st.error("Username orang tua sudah ada di tabel users.")
                conn.close()
                return

        cur.execute("""
        INSERT INTO users (username, password, role, nama, status_akun)
        VALUES (?, ?, 'siswa', ?, 'aktif')
        """, (
            data["username_siswa"],
            data["password_siswa"],
            data["nama_siswa"]
        ))

        user_siswa_id = cur.lastrowid

        cur.execute("""
        INSERT INTO siswa (user_id, nama, kelas, sekolah, email)
        VALUES (?, ?, ?, ?, ?)
        """, (
            user_siswa_id,
            data["nama_siswa"],
            data["kelas"],
            data["sekolah"],
            data["email_siswa"]
        ))

        siswa_id = cur.lastrowid

        cur.execute("""
        INSERT INTO users (username, password, role, nama, status_akun)
        VALUES (?, ?, 'ortu', ?, 'aktif')
        """, (
            data["username_ortu"],
            data["password_ortu"],
            data["nama_ortu"]
        ))

        user_ortu_id = cur.lastrowid

        cur.execute("""
        INSERT INTO ortu (user_id, siswa_id, nama, hubungan, email, no_hp)
        VALUES (?, ?, ?, ?, ?, ?)
        """, (
            user_ortu_id,
            siswa_id,
            data["nama_ortu"],
            data["hubungan"],
            data["email_ortu"],
            data["no_hp"]
        ))

        cur.execute("""
        UPDATE pendaftaran_ortu_siswa
        SET status = 'diterima',
            catatan_admin = ?,
            updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """, (catatan_admin, selected_id))

        conn.commit()
        st.success("Registrasi diterima. Akun siswa dan orang tua berhasil dibuat.")
        st.rerun()

    except Exception as e:
        conn.rollback()
        st.error("Gagal menerima registrasi.")
        st.code(str(e))


def proses_tolak_registrasi(conn, selected_id, catatan_admin):
    cur = conn.cursor()

    try:
        cur.execute("""
        UPDATE pendaftaran_ortu_siswa
        SET status = 'ditolak',
            catatan_admin = ?,
            updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """, (catatan_admin, selected_id))

        conn.commit()
        st.warning("Registrasi ditolak.")
        st.rerun()

    except Exception as e:
        conn.rollback()
        st.error("Gagal menolak registrasi.")
        st.code(str(e))


# =================
# NOTIFIKASI STATUS REGISTRASI
# ======================
def page_cek_status_registrasi():
    st.title("🔎 Cek Status Registrasi")

    st.info("Masukkan username orang tua atau username siswa yang digunakan saat registrasi.")

    username = st.text_input("Username registrasi")

    if st.button("Cek Status"):
        if not username:
            st.warning("Username wajib diisi.")
            return

        conn = connect_db()

        df = pd.read_sql_query("""
        SELECT 
            id,
            nama_siswa,
            username_siswa,
            nama_ortu,
            username_ortu,
            status,
            catatan_admin,
            created_at,
            updated_at
        FROM pendaftaran_ortu_siswa
        WHERE username_siswa = ? OR username_ortu = ?
        ORDER BY created_at DESC
        LIMIT 1
        """, conn, params=(username, username))

        conn.close()

        if df.empty:
            st.error("Data registrasi tidak ditemukan.")
            return

        data = df.iloc[0]

        if data["status"] == "pending":
            st.warning("Status registrasi: PENDING. Registrasi masih menunggu verifikasi admin.")
        elif data["status"] == "diterima":
            st.success("Status registrasi: DITERIMA. Akun siswa dan orang tua sudah bisa digunakan untuk login.")
        elif data["status"] == "ditolak":
            st.error("Status registrasi: DITOLAK.")

        st.write(f"Nama siswa: **{data['nama_siswa']}**")
        st.write(f"Nama orang tua: **{data['nama_ortu']}**")
        st.write(f"Tanggal registrasi: **{data['created_at']}**")

        if data["catatan_admin"]:
            st.info(f"Catatan admin: {data['catatan_admin']}")


# =================
# CRUD SISWA
# =================
def page_crud_siswa():
    st.title("👨‍🎓 Data Siswa")

    conn = connect_db()

    df = pd.read_sql_query("""
    SELECT 
        siswa.id, 
        users.username, 
        users.status_akun,
        siswa.nama, 
        siswa.kelas, 
        siswa.sekolah, 
        siswa.email
    FROM siswa
    JOIN users ON siswa.user_id = users.id
    ORDER BY siswa.id DESC
    """, conn)

    st.dataframe(df, use_container_width=True)

    st.divider()
    st.subheader("Tambah Siswa")

    with st.form("form_tambah_siswa"):
        nama = st.text_input("Nama Siswa")
        kelas = st.text_input("Kelas")
        sekolah = st.text_input("Sekolah")
        email = st.text_input("Email")
        username = st.text_input("Username Login")
        password = st.text_input("Password Login", type="password")

        submit = st.form_submit_button("Simpan Siswa")

        if submit:
            cur = conn.cursor()

            try:
                cur.execute("""
                INSERT INTO users (username, password, role, nama)
                VALUES (?, ?, ?, ?)
                """, (username, password, "siswa", nama))

                user_id = cur.lastrowid

                cur.execute("""
                INSERT INTO siswa (user_id, nama, kelas, sekolah, email)
                VALUES (?, ?, ?, ?, ?)
                """, (user_id, nama, kelas, sekolah, email))

                conn.commit()
                st.success("Data siswa berhasil ditambahkan.")
                st.rerun()

            except Exception as e:
                conn.rollback()
                st.error("Gagal menambahkan siswa.")
                st.code(str(e))

    st.divider()
    st.subheader("Edit / Hapus Siswa")

    if not df.empty:
        selected_id = st.selectbox("Pilih ID Siswa", df["id"].tolist())
        siswa_lama = df[df["id"] == selected_id].iloc[0]

        with st.form("form_edit_siswa"):
            nama = st.text_input("Nama Siswa", siswa_lama["nama"])
            kelas = st.text_input("Kelas", siswa_lama["kelas"])
            sekolah = st.text_input("Sekolah", siswa_lama["sekolah"])
            email = st.text_input("Email", siswa_lama["email"])
            username = st.text_input("Username Login", siswa_lama["username"])
            status_akun = st.selectbox(
                "Status Akun",
                ["aktif", "nonaktif"],
                index=["aktif", "nonaktif"].index(siswa_lama["status_akun"])
                if siswa_lama["status_akun"] in ["aktif", "nonaktif"] else 0
            )

            col1, col2 = st.columns(2)

            update_btn = col1.form_submit_button("Update")
            delete_btn = col2.form_submit_button("Hapus")

            cur = conn.cursor()

            if update_btn:
                cur.execute("""
                UPDATE siswa
                SET nama=?, kelas=?, sekolah=?, email=?
                WHERE id=?
                """, (nama, kelas, sekolah, email, selected_id))

                cur.execute("""
                UPDATE users
                SET nama=?, username=?, status_akun=?
                WHERE id = (
                    SELECT user_id FROM siswa WHERE id=?
                )
                """, (nama, username, status_akun, selected_id))

                conn.commit()
                st.success("Data siswa berhasil diupdate.")
                st.rerun()

            if delete_btn:
                cur.execute("""
                DELETE FROM users
                WHERE id = (
                    SELECT user_id FROM siswa WHERE id=?
                )
                """, (selected_id,))

                cur.execute("DELETE FROM siswa WHERE id=?", (selected_id,))

                conn.commit()
                st.success("Data siswa berhasil dihapus.")
                st.rerun()

    conn.close()


# =====================
# CRUD ORTU
# =====================
def page_crud_ortu():
    st.title("👨‍👩‍👧 Data Orang Tua")

    conn = connect_db()

    df_siswa = pd.read_sql_query("""
    SELECT id, nama, kelas, sekolah
    FROM siswa
    ORDER BY nama ASC
    """, conn)

    df = pd.read_sql_query("""
    SELECT 
        ortu.id,
        users.username,
        users.status_akun,
        ortu.nama,
        ortu.hubungan,
        ortu.email,
        ortu.no_hp,
        password,
        siswa.nama AS nama_siswa
    FROM ortu
    JOIN users ON ortu.user_id = users.id
    JOIN siswa ON ortu.siswa_id = siswa.id
    ORDER BY ortu.id DESC
    """, conn)

    st.dataframe(df, use_container_width=True)

    st.divider()
    st.subheader("Tambah Orang Tua")

    if df_siswa.empty:
        st.warning("Belum ada data siswa. Tambahkan siswa dulu.")
    else:
        siswa_options = {
            f"{row['id']} - {row['nama']} ({row['kelas']})": row["id"]
            for _, row in df_siswa.iterrows()
        }

        with st.form("form_tambah_ortu"):
            nama = st.text_input("Nama Orang Tua")
            hubungan = st.selectbox("Hubungan", ["Ayah", "Ibu", "Wali"])
            email = st.text_input("Email")
            no_hp = st.text_input("No HP")
            siswa_label = st.selectbox("Anak/Siswa", list(siswa_options.keys()))
            username = st.text_input("Username Login")
            password = st.text_input("Password Login", type="password")

            submit = st.form_submit_button("Simpan Orang Tua")

            if submit:
                cur = conn.cursor()

                try:
                    cur.execute("""
                    INSERT INTO users (username, password, role, nama)
                    VALUES (?, ?, ?, ?)
                    """, (username, password, "ortu", nama))

                    user_id = cur.lastrowid
                    siswa_id = siswa_options[siswa_label]

                    cur.execute("""
                    INSERT INTO ortu (user_id, siswa_id, nama, hubungan, email, no_hp)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """, (user_id, siswa_id, nama, hubungan, email, no_hp))

                    conn.commit()
                    st.success("Data orang tua berhasil ditambahkan.")
                    st.rerun()

                except Exception as e:
                    conn.rollback()
                    st.error("Gagal menambahkan orang tua.")
                    st.code(str(e))

    st.divider()
    st.subheader("Edit / Hapus Orang Tua")

    if not df.empty:
        selected_id = st.selectbox("Pilih ID Orang Tua", df["id"].tolist())
        ortu_lama = df[df["id"] == selected_id].iloc[0]

        siswa_options = {
            f"{row['id']} - {row['nama']} ({row['kelas']})": row["id"]
            for _, row in df_siswa.iterrows()
        }

        with st.form("form_edit_ortu"):
            nama = st.text_input("Nama Orang Tua", ortu_lama["nama"])
            hubungan = st.selectbox(
                "Hubungan",
                ["Ayah", "Ibu", "Wali"],
                index=["Ayah", "Ibu", "Wali"].index(ortu_lama["hubungan"])
                if ortu_lama["hubungan"] in ["Ayah", "Ibu", "Wali"] else 0
            )
            email = st.text_input("Email", ortu_lama["email"])
            no_hp = st.text_input("No HP", ortu_lama["no_hp"])
            username = st.text_input("Username Login", ortu_lama["username"])
            siswa_label = st.selectbox("Anak/Siswa", list(siswa_options.keys()))
            status_akun = st.selectbox(
                "Status Akun",
                ["aktif", "nonaktif"],
                index=["aktif", "nonaktif"].index(ortu_lama["status_akun"])
                if ortu_lama["status_akun"] in ["aktif", "nonaktif"] else 0
            )

            col1, col2 = st.columns(2)

            update_btn = col1.form_submit_button("Update")
            delete_btn = col2.form_submit_button("Hapus")

            cur = conn.cursor()

            if update_btn:
                siswa_id = siswa_options[siswa_label]

                cur.execute("""
                UPDATE ortu
                SET nama=?, hubungan=?, email=?, no_hp=?, siswa_id=?
                WHERE id=?
                """, (nama, hubungan, email, no_hp, siswa_id, selected_id))

                cur.execute("""
                UPDATE users
                SET nama=?, username=?, status_akun=?
                WHERE id = (
                    SELECT user_id FROM ortu WHERE id=?
                )
                """, (nama, username, status_akun, selected_id))

                conn.commit()
                st.success("Data orang tua berhasil diupdate.")
                st.rerun()

            if delete_btn:
                cur.execute("""
                DELETE FROM users
                WHERE id = (
                    SELECT user_id FROM ortu WHERE id=?
                )
                """, (selected_id,))

                cur.execute("DELETE FROM ortu WHERE id=?", (selected_id,))

                conn.commit()
                st.success("Data orang tua berhasil dihapus.")
                st.rerun()

    conn.close()


def logout():
    visitor_session_id = st.session_state.get("visitor_session_id")

    st.session_state.clear()

    if visitor_session_id:
        st.session_state["visitor_session_id"] = visitor_session_id

    st.rerun()


def go_to(menu_name):
    st.session_state["menu"] = menu_name
    st.rerun()


# =========================
# LOGIN PAGE
# =========================

def page_login():
    st.title("📘 SMART TKA Digital")
    st.subheader("Login")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        user = login(username, password)

        if user == "nonaktif":
            st.error("Akun kamu sedang nonaktif. Silakan hubungi admin.")
        elif user:
            st.session_state["username"] = user[0]
            st.session_state["role"] = user[1]
            st.session_state["nama"] = user[2]
            st.session_state["menu"] = "Dashboard"
            catat_login_pengunjung(user[0], user[1])
            st.success("Login berhasil")
            st.rerun()
        else:
            st.error("Username atau password salah")

    st.divider()

    st.info(
        "Belum punya akun? Daftar akun penuh Smart TKA untuk menyimpan progress, "
        "mendapat try out adaptif, learning planner, jadwal belajar, dan dashboard orang tua."
    )
    
    if st.button("Daftar Berbayar Rp100.000"):
        st.session_state["halaman_awal"] = "register_bayar"
        st.rerun()
# =========================
# DASHBOARD
# =========================

def card(title, desc, button_label, target_menu, key):
    with st.container(border=True):
        col1, col2 = st.columns([2, 1])

        with col1:
            st.subheader(title)
            st.caption(desc)

        with col2:
            st.write("")
            if st.button(button_label, key=key, use_container_width=True):
                go_to(target_menu)


def page_dashboard():
    role = st.session_state["role"]

    st.title("🏠 Dashboard")
    st.write(f"Halo, **{st.session_state['nama']}**")
    # st.write(f"Role: **{role}**")

    # st.divider()

    if role == "admin":
        # st.subheader("Fitur Admin")

        col1, col2 = st.columns(2)

        with col1:
            card(
                "🗂️ Bank Soal",
                "CRUD data soal TKA.",
                "Buka Bank Soal",
                "Bank Soal",
                "admin_bank_soal"
            )

        with col2:
            card(
                "📊 Hasil Try Out",
                "Lihat rekap hasil try out seluruh siswa.",
                "Buka Hasil",
                "Hasil Try Out",
                "admin_hasil_tryout"
            )

        col3, col4 = st.columns(2)
        with col3:
            card(
                "📚 Upload Materi",
                "Upload dokumen materi sebagai sumber jawaban AI Tutor.",
                "Upload",
                "Upload Materi",
                "admin_upload_materi"
            )

        with col4:
            card(
                "🤖 Generate Soal AI",
                "Buat soal otomatis dari materi yang sudah diupload.",
                "Generate",
                "Generate Soal AI",
                "admin_generate_soal"
            )

        col5, col6 = st.columns(2)

        with col5:
            card(
                "💬 Uji Coba AI Tutor",
                "Uji coba fitur tutor berbasis LLM pada aplikasi.",
                "Buka AI Tutor",
                "AI Tutor",
                "admin_tutor"
            )

        with col6:
            card(
                "👨‍🎓 Data Siswa",
                "Kelola data siswa dan akun login siswa.",
                "Buka",
                "Data Siswa",
                "admin_data_siswa"
            )

        col7, col8 = st.columns(2)
        with col7:
            card(
                "👨‍👩‍👧 Data Orang Tua",
                "Kelola data orang tua dan hubungkan dengan siswa.",
                "Buka",
                "Data Orang Tua",
                "admin_data_ortu"
            )

        with col8:
            card(
                "✅ Verifikasi Register",
                "Terima atau tolak pendaftaran orang tua dan siswa.",
                "Verifikasi",
                "Verifikasi Register",
                "admin_verifikasi_register"
            )

    elif role == "siswa":
        st.subheader("Fitur Siswa")
        tampilkan_notifikasi_belajar()
        st.divider()
        col1, col2 = st.columns(2)

        with col1:
            card(
                "📝 Simulasi TKA Digital",
                "Kerjakan try out TKA dan dapatkan nilai otomatis.",
                "Mulai",
                "Simulasi TKA",
                "siswa_tryout"
            )

        with col2:
            card(
                "🎯 Learning Planner",
                "Buat rencana belajar berdasarkan hasil try out terakhir.",
                "Buka Planner",
                "Learning Planner",
                "siswa_learning_planner"
            )

        col3, col4 = st.columns(2)

        with col3:
            card(
                "📈 Progress Belajar",
                "Lihat perkembangan hasil try out kamu.",
                "Lihat Progress",
                "Progress Saya",
                "siswa_progress"
            )

        with col4:
            card(
                "💬 AI Tutor",
                "Tanyakan soal atau materi yang belum dipahami.",
                "Buka",
                "AI Tutor",
                "siswa_tutor"
            )

        col5, col6 = st.columns(2)

        with col5:
            card(
                "🧭 Adaptive Learning",
                "Rekomendasi belajar otomatis berdasarkan topik yang sering salah.",
                "Buka",
                "Adaptive Learning",
                "siswa_adaptive_learning"
            )

        with col6:
            card(
                "📅 Jadwal Belajar",
                "Buat jadwal belajar otomatis berdasarkan hasil try out.",
                "Buka",
                "Jadwal Belajar",
                "siswa_jadwal_belajar"
            )

    elif role == "ortu":
        st.subheader("Fitur Orang Tua")

        col1, col2 = st.columns(2)

        with col1:
            card(
                "👨‍👩‍👧 Dashboard Progress Anak",
                "Pantau nilai, perkembangan, dan rekomendasi belajar anak.",
                "Lihat Progress Anak",
                "Dashboard Orang Tua",
                "ortu_progress"
            )

        with col2:
            st.info("📌 Catatan")
            st.write("Orang tua hanya dapat melihat progress belajar dan hasil try out.")


# =========================
# TAMBAH FUNGSI BACA FILE
# =========================
def baca_file_upload(uploaded_file):
    filename = uploaded_file.name.lower()

    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    elif filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif filename.endswith(".docx"):
        document = docx.Document(uploaded_file)
        text = ""
        for para in document.paragraphs:
            text += para.text + "\n"
        return text

    else:
        return ""


# =========================
# TAMBAH FUNGSI BACA FILE
# =========================
def cari_materi_relevan(pertanyaan, limit=5):
    conn = connect_db()

    df = pd.read_sql_query("""
    SELECT
        materi_chunk.id,
        materi_chunk.judul,
        materi_chunk.isi_chunk,
        materi_chunk.urutan
    FROM materi_chunk
    ORDER BY materi_chunk.id ASC
    """, conn)

    conn.close()

    if df.empty:
        return ""

    # Bersihkan pertanyaan
    teks_tanya = pertanyaan.lower()

    for simbol in ["?", ".", ",", "!", ":", ";", "(", ")"]:
        teks_tanya = teks_tanya.replace(simbol, " ")

    stopwords = {
        "apa", "itu", "yang", "dan", "di", "ke",
        "dari", "adalah", "bagaimana", "jelaskan",
        "cara", "dengan", "untuk", "pada",
        "kah", "nya", "saya", "aku"
    }

    kata_kunci = []

    for kata in teks_tanya.split():
        if kata not in stopwords:
            if len(kata) > 2:
                kata_kunci.append(kata)

    hasil = []

    for _, row in df.iterrows():

        isi = str(row["isi_chunk"]).lower()

        skor = 0

        for kata in kata_kunci:

            # bobot tinggi jika cocok persis
            skor += isi.count(kata) * 3

            # bonus jika kata muncul di judul
            if kata in str(row["judul"]).lower():
                skor += 10

        if skor > 0:
            hasil.append({
                "judul": row["judul"],
                "urutan": row["urutan"],
                "isi": row["isi_chunk"],
                "skor": skor
            })

    hasil = sorted(
        hasil,
        key=lambda x: x["skor"],
        reverse=True
    )

    konteks = ""

    for item in hasil[:limit]:
        konteks += f"""
        Judul Materi: {item['judul']}
        Bagian: {item['urutan']}

        {item['isi']}

        ====================================
        """

    return konteks


# =========================
# HALAMAN UPLOAD MATERI
# =========================
def page_materi():
    st.title("📚 Upload Materi Pembelajaran")

    conn = connect_db()

    st.subheader("Daftar Materi")

    df = pd.read_sql_query("""
    SELECT id, judul, filename, created_at
    FROM materi
    ORDER BY created_at DESC
    """, conn)

    st.dataframe(df, use_container_width=True)

    st.divider()

    st.subheader("Upload Materi Baru")

    judul = st.text_input("Judul Materi")
    uploaded_file = st.file_uploader(
        "Upload file materi",
        type=["txt", "pdf", "docx"]
    )

    if st.button("Simpan Materi"):
        if not judul or uploaded_file is None:
            st.warning("Judul dan file wajib diisi.")
        else:
            isi_teks = baca_file_upload(uploaded_file)

            if isi_teks.strip() == "":
                st.error("Isi file tidak terbaca.")
            else:
                cur = conn.cursor()
                cur.execute("""
                INSERT INTO materi (judul, filename, isi_teks)
                VALUES (?, ?, ?)
                """, (judul, uploaded_file.name, isi_teks))

                materi_id = cur.lastrowid

                chunks = pecah_teks_jadi_chunk(isi_teks)

                for i, chunk in enumerate(chunks):
                    cur.execute("""
                    INSERT INTO materi_chunk (materi_id, judul, isi_chunk, urutan)
                    VALUES (?, ?, ?, ?)
                    """, (materi_id, judul, chunk, i + 1))

                conn.commit()
                st.success("Materi berhasil disimpan.")
                st.rerun()

    st.divider()

    st.subheader("Hapus Materi")

    if not df.empty:
        selected_id = st.selectbox("Pilih ID Materi", df["id"].tolist())

        if st.button("Hapus Materi"):
            cur = conn.cursor()
            cur.execute("DELETE FROM materi WHERE id=?", (selected_id,))
            conn.commit()
            st.success("Materi berhasil dihapus.")
            st.rerun()

    conn.close()


# =========================
# PECAH ISI FILE MATERI
# =========================
def pecah_teks_jadi_chunk(teks, ukuran=1500, overlap=300):
    chunks = []
    start = 0

    while start < len(teks):
        end = start + ukuran
        chunk = teks[start:end]

        if chunk.strip():
            chunks.append(chunk.strip())

        start += ukuran - overlap

    return chunks


# =========================
# CRUD BANK SOAL
# =========================

def page_bank_soal():
    st.title("🗂️ CRUD Bank Soal")

    conn = connect_db()

    st.subheader("Data Bank Soal")
    df = pd.read_sql_query("SELECT * FROM bank_soal", conn)
    st.dataframe(df, use_container_width=True)

    st.divider()

    st.subheader("Tambah Soal Manual")

    with st.form("form_tambah_soal"):
        mapel = st.selectbox("Mapel", MAPEL_OPTIONS)
        topik = st.text_input("Topik")
        level = st.selectbox("Level Kesulitan", ["Mudah", "Sedang", "Sulit"])
        pertanyaan = st.text_area("Pertanyaan")
        opsi_a = st.text_input("Opsi A")
        opsi_b = st.text_input("Opsi B")
        opsi_c = st.text_input("Opsi C")
        opsi_d = st.text_input("Opsi D")
        jawaban = st.selectbox("Jawaban Benar", ["A", "B", "C", "D"])
        pembahasan = st.text_area("Pembahasan")

        submitted = st.form_submit_button("Simpan Soal")

        if submitted:
            cur = conn.cursor()
            cur.execute("""
            INSERT INTO bank_soal
            (mapel, topik, level, pertanyaan, opsi_a, opsi_b, opsi_c, opsi_d, jawaban, pembahasan)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (mapel, topik, level, pertanyaan, opsi_a, opsi_b, opsi_c, opsi_d, jawaban, pembahasan))
            conn.commit()
            st.success("Soal berhasil ditambahkan")
            st.rerun()

    st.divider()

    st.subheader("Edit / Hapus Soal")

    if not df.empty:
        selected_id = st.selectbox("Pilih ID Soal", df["id"].tolist())
        soal = df[df["id"] == selected_id].iloc[0]

        with st.form("form_edit_soal"):
            mapel = st.selectbox(
                "Mapel",
                MAPEL_OPTIONS,
                index=MAPEL_OPTIONS.index(soal["mapel"]) if soal["mapel"] in MAPEL_OPTIONS else 0
            )
            topik = st.text_input("Topik", soal["topik"])
            level = st.selectbox(
                "Level Kesulitan",
                ["Mudah", "Sedang", "Sulit"],
                index=["Mudah", "Sedang", "Sulit"].index(soal["level"]) if soal["level"] in ["Mudah", "Sedang",
                                                                                             "Sulit"] else 0
            )
            pertanyaan = st.text_area("Pertanyaan", soal["pertanyaan"])
            opsi_a = st.text_input("Opsi A", soal["opsi_a"])
            opsi_b = st.text_input("Opsi B", soal["opsi_b"])
            opsi_c = st.text_input("Opsi C", soal["opsi_c"])
            opsi_d = st.text_input("Opsi D", soal["opsi_d"])
            jawaban = st.selectbox(
                "Jawaban Benar",
                ["A", "B", "C", "D"],
                index=["A", "B", "C", "D"].index(soal["jawaban"])
            )
            pembahasan = st.text_area("Pembahasan", soal["pembahasan"])

            col1, col2 = st.columns(2)

            update_btn = col1.form_submit_button("Update")
            delete_btn = col2.form_submit_button("Hapus")

            cur = conn.cursor()

            if update_btn:
                cur.execute("""
                UPDATE bank_soal
                SET mapel=?, topik=?, level=?, pertanyaan=?, opsi_a=?, opsi_b=?, opsi_c=?, opsi_d=?, jawaban=?, pembahasan=?
                WHERE id=?
                """, (
                    mapel, topik, level, pertanyaan, opsi_a, opsi_b, opsi_c, opsi_d, jawaban, pembahasan, selected_id))
                conn.commit()
                st.success("Soal berhasil diupdate")
                st.rerun()

            if delete_btn:
                cur.execute("DELETE FROM bank_soal WHERE id=?", (selected_id,))
                conn.commit()
                st.success("Soal berhasil dihapus")
                st.rerun()

    conn.close()


# =========================
# SIMULASI TKA
# =========================

def page_tryout():
    st.title("📝 Simulasi TKA Digital Adaptif")

    conn = connect_db()

    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        st.error("Data siswa tidak ditemukan. Pastikan akun siswa sudah terhubung dengan tabel siswa.")
        conn.close()
        return

    mapel_dipilih = st.selectbox(
        "Pilih Mapel TKA",
        MAPEL_OPTIONS
    )

    jumlah_soal = st.number_input(
        "Jumlah soal",
        min_value=1,
        max_value=30,
        value=10
    )

    st.info(
        "Soal akan dipilih otomatis oleh sistem berdasarkan riwayat kemampuan kamu sebelumnya."
    )

    paket_key = f"paket_adaptif_{siswa_id}_{mapel_dipilih}_{jumlah_soal}"

    if st.button("Buat Paket Soal Adaptif"):
        df_paket, info_adaptif = ambil_soal_adaptif(
            siswa_id=siswa_id,
            mapel=mapel_dipilih,
            jumlah_soal=jumlah_soal
        )

        if df_paket.empty:
            st.warning(f"Belum ada soal untuk mapel {mapel_dipilih}.")
            conn.close()
            return

        st.session_state[paket_key] = df_paket["id"].tolist()
        st.session_state[f"{paket_key}_info"] = info_adaptif
        st.success("Paket soal adaptif berhasil dibuat.")

    if paket_key not in st.session_state:
        st.warning("Klik tombol **Buat Paket Soal Adaptif** dulu untuk memulai.")
        conn.close()
        return

    soal_ids = st.session_state[paket_key]

    placeholders = ",".join(["?"] * len(soal_ids))

    df = pd.read_sql_query(f"""
    SELECT *
    FROM bank_soal
    WHERE id IN ({placeholders})
    """, conn, params=soal_ids)

    # Supaya urutan soal tetap sama seperti paket adaptif yang dibuat
    df["urutan_paket"] = df["id"].apply(lambda x: soal_ids.index(x))
    df = df.sort_values("urutan_paket")

    info_adaptif = st.session_state.get(f"{paket_key}_info", pd.DataFrame())

    if not info_adaptif.empty:
        with st.expander("🧭 Informasi Adaptasi Level"):
            st.dataframe(info_adaptif, use_container_width=True)

    st.info(f"Mapel: {mapel_dipilih} | Jumlah soal: {len(df)}")

    jawaban_user = {}

    with st.form("form_tryout"):
        for _, row in df.iterrows():
            st.markdown(f"### {row['pertanyaan']}")
            st.caption(f"Mapel: {row['mapel']} | Topik: {row['topik']} | Level: {row['level']}")

            pilihan = {
                "A": row["opsi_a"],
                "B": row["opsi_b"],
                "C": row["opsi_c"],
                "D": row["opsi_d"],
            }

            jawaban_user[row["id"]] = st.radio(
                "Pilih jawaban",
                options=["A", "B", "C", "D"],
                index=None,
                format_func=lambda x: f"{x}. {pilihan[x]}",
                key=f"soal_{row['id']}"
            )

        submit = st.form_submit_button("Submit Try Out")

    if submit:
        benar = 0
        total = len(df)

        hasil_mapel = {}
        topik_benar = []
        topik_salah = []

        for _, row in df.iterrows():
            mapel = row["mapel"]
            topik = row["topik"]

            if mapel not in hasil_mapel:
                hasil_mapel[mapel] = {
                    "benar": 0,
                    "total": 0
                }

            hasil_mapel[mapel]["total"] += 1

            if jawaban_user[row["id"]] == row["jawaban"]:
                benar += 1
                hasil_mapel[mapel]["benar"] += 1
                topik_benar.append(topik)
            else:
                topik_salah.append(topik)

        nilai = round((benar / total) * 100, 2)

        ringkasan_mapel = ""
        for mapel, data in hasil_mapel.items():
            nilai_mapel = round((data["benar"] / data["total"]) * 100, 2)
            ringkasan_mapel += f"\n- {mapel}: {nilai_mapel} ({data['benar']} dari {data['total']} benar)"

        topik_benar_unik = list(set(topik_benar))
        topik_salah_unik = list(set(topik_salah))

        prompt = f"""
        Analisis hasil try out TKA siswa berikut.

        Mapel: {mapel_dipilih}
        Nilai total: {nilai}
        Jawaban benar: {benar} dari {total}

        Nilai per mapel:
        {ringkasan_mapel}

        Topik yang dikuasai:
        {topik_benar_unik}

        Topik yang masih salah:
        {topik_salah_unik}

        Buat analisis dalam format:
        1. Ringkasan kemampuan siswa
        2. Kekuatan siswa
        3. Kelemahan siswa
        4. Rekomendasi materi yang harus dipelajari
        5. Rencana belajar 7 hari
        6. Motivasi singkat

        Gunakan bahasa Indonesia yang ramah, jelas, dan memotivasi.
        """

        rekomendasi = ask_llm(prompt)

        # cur = conn.cursor()
        siswa_id = get_siswa_id_by_username(st.session_state["username"])

        if siswa_id is None:
            st.error("Data siswa tidak ditemukan. Pastikan akun siswa sudah terhubung dengan tabel siswa.")
            conn.close()
            return

        try:
            cur = conn.cursor()

            cur.execute("""
            INSERT INTO hasil_tryout (siswa_id, mapel, nilai, benar, total, rekomendasi)
            VALUES (?, ?, ?, ?, ?, ?)
            """, (siswa_id, mapel_dipilih, nilai, benar, total, rekomendasi))

            hasil_tryout_id = cur.lastrowid

            for _, row in df.iterrows():
                jawaban_siswa = jawaban_user[row["id"]]
                jawaban_benar = row["jawaban"]
                is_benar = 1 if jawaban_siswa == jawaban_benar else 0

                cur.execute("""
                INSERT INTO hasil_tryout_detail
                (hasil_tryout_id, siswa_id, soal_id, mapel, topik, level, jawaban_siswa, jawaban_benar, is_benar)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    hasil_tryout_id,
                    siswa_id,
                    row["id"],
                    mapel_dipilih,
                    row["topik"],
                    row["level"],
                    jawaban_siswa,
                    jawaban_benar,
                    is_benar
                ))

            conn.commit()

            st.success("Hasil try out berhasil disimpan ke database.")

        except Exception as e:
            conn.rollback()
            st.error("Gagal menyimpan hasil try out.")
            st.code(str(e))
            conn.close()
            return

        st.success(f"Nilai kamu: {nilai}")
        st.write(f"Benar: {benar} dari {total}")

        st.divider()

        st.subheader("🧾 Review Soal & Pembahasan")
        
        for nomor, (_, row) in enumerate(df.iterrows(), start=1):
            jawaban_siswa = jawaban_user[row["id"]]
            jawaban_benar = row["jawaban"]
        
            pilihan = {
                "A": row["opsi_a"],
                "B": row["opsi_b"],
                "C": row["opsi_c"],
                "D": row["opsi_d"],
            }
        
            is_benar = jawaban_siswa == jawaban_benar
            status = "✅ Benar" if is_benar else "❌ Salah"
        
            topik = row["topik"] if pd.notna(row["topik"]) else "-"
            level = row["level"] if pd.notna(row["level"]) else "-"
        
            with st.expander(
                f"Soal {nomor} | {status} | Topik: {topik} | Level: {level}",
                expanded=False
            ):
                st.markdown(f"**Pertanyaan:** {row['pertanyaan']}")
        
                st.write("**Pilihan Jawaban:**")
        
                for kode, teks_opsi in pilihan.items():
                    label = ""
        
                    if kode == jawaban_benar and kode == jawaban_siswa:
                        label = " ✅ Jawaban kamu & jawaban benar"
                    elif kode == jawaban_benar:
                        label = " ✅ Jawaban benar"
                    elif kode == jawaban_siswa:
                        label = " ❌ Jawaban kamu"
        
                    st.write(f"{kode}. {teks_opsi}{label}")
        
                if is_benar:
                    st.success(
                        f"Jawaban kamu benar: {jawaban_siswa}. {pilihan.get(jawaban_siswa, '-')}"
                    )
                else:
                    st.error(
                        f"Jawaban kamu: {jawaban_siswa}. {pilihan.get(jawaban_siswa, '-')}"
                    )
                    st.success(
                        f"Jawaban benar: {jawaban_benar}. {pilihan.get(jawaban_benar, '-')}"
                    )
        
                pembahasan = row["pembahasan"]
        
                if pd.isna(pembahasan) or str(pembahasan).strip() == "":
                    pembahasan = "Pembahasan belum tersedia."
        
                st.markdown("**Pembahasan:**")
                st.info(pembahasan)

        st.subheader("📊 Analisis Per Mapel")
        for mapel, data in hasil_mapel.items():
            nilai_mapel = round((data["benar"] / data["total"]) * 100, 2)
            st.write(f"**{mapel}**: {nilai_mapel} — {data['benar']} dari {data['total']} benar")

        st.subheader("✅ Topik Dikuasai")
        st.write(topik_benar_unik if topik_benar_unik else "-")

        st.subheader("⚠️ Topik Perlu Dipelajari")
        st.write(topik_salah_unik if topik_salah_unik else "-")

        st.subheader("🤖 Diagnosis AI")
        st.write(rekomendasi)

    conn.close()


# =========================
# PAGE ADAPTIVE LEARNING
# =========================
def page_adaptive_learning():
    st.title("🧭 Adaptive Learning")

    conn = connect_db()

    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        st.error("Data siswa tidak ditemukan.")
        conn.close()
        return

    mapel_dipilih = st.selectbox(
        "Pilih mapel",
        MAPEL_OPTIONS
    )

    df = pd.read_sql_query("""
    SELECT 
        topik,
        level,
        COUNT(*) AS total_dikerjakan,
        SUM(CASE WHEN is_benar = 1 THEN 1 ELSE 0 END) AS total_benar,
        SUM(CASE WHEN is_benar = 0 THEN 1 ELSE 0 END) AS total_salah,
        ROUND(AVG(is_benar) * 100, 2) AS akurasi
    FROM hasil_tryout_detail
    WHERE siswa_id = ? AND mapel = ?
    GROUP BY topik, level
    ORDER BY akurasi ASC, total_salah DESC
    """, conn, params=(siswa_id, mapel_dipilih))

    conn.close()

    if df.empty:
        st.info(f"Belum ada data detail try out untuk mapel {mapel_dipilih}. Kerjakan Simulasi TKA dulu.")
        return

    st.subheader("📊 Analisis Topik")

    st.dataframe(df, use_container_width=True)

    st.divider()

    topik_lemah = df[df["akurasi"] < 80]

    if topik_lemah.empty:
        st.success("Mantap! Belum ada topik yang terlihat lemah. Pertahankan latihan dan coba soal level lebih tinggi.")
        return

    st.subheader("⚠️ Topik Prioritas")

    for _, row in topik_lemah.iterrows():
        with st.container(border=True):
            st.write(f"**Topik:** {row['topik']}")
            st.write(f"**Level:** {row['level']}")
            st.write(f"**Akurasi:** {row['akurasi']}%")
            st.write(f"**Salah:** {row['total_salah']} dari {row['total_dikerjakan']} soal")

    st.divider()

    if st.button("Buat Rekomendasi Adaptive Learning"):
        daftar_topik = ""

        for _, row in topik_lemah.head(5).iterrows():
            daftar_topik += f"""
            - Topik: {row['topik']}
              Level: {row['level']}
              Akurasi: {row['akurasi']}%
              Salah: {row['total_salah']} dari {row['total_dikerjakan']} soal
            """

        konteks_materi = ""

        for _, row in topik_lemah.head(3).iterrows():
            query_materi = f"{mapel_dipilih} {row['topik']}"
            potongan_materi = cari_materi_relevan(query_materi, limit=2)

            if potongan_materi.strip():
                konteks_materi += f"""
                Materi terkait topik {row['topik']}:
                {potongan_materi[:2000]}
                """

        prompt = f"""
        Kamu adalah sistem Adaptive Learning untuk siswa TKA.

        Mapel: {mapel_dipilih}

        Data topik yang perlu diperbaiki:
        {daftar_topik}

        Materi pendukung dari dokumen yang diupload:
        {konteks_materi}

        Buat rekomendasi belajar adaptif dengan format:

        1. Ringkasan kondisi siswa
        2. Urutan prioritas belajar
        3. Materi yang harus dipelajari ulang
        4. Jenis latihan soal yang disarankan
        5. Rencana belajar 5 hari
        6. Pertanyaan yang sebaiknya siswa tanyakan ke AI Tutor
        7. Motivasi singkat

        Aturan:
        - Fokus pada topik dengan akurasi paling rendah.
        - Jangan menyarankan materi di luar mapel yang dipilih.
        - Gunakan bahasa Indonesia yang ramah dan mudah dipahami.
        """

        with st.spinner("AI sedang membuat rekomendasi belajar adaptif..."):
            hasil = ask_llm(prompt)

        st.subheader("🧭 Rekomendasi Adaptive Learning")
        st.write(hasil)


# =========================
# AI TUTOR
# =========================

def page_ai_tutor():
    st.title("💬 AI Tutor Berbasis Materi")

    pertanyaan = st.text_area("Tanyakan soal atau materi TKA")

    if st.button("Tanya AI"):
        if pertanyaan.strip() == "":
            st.warning("Pertanyaan belum diisi.")
        else:
            konteks = cari_materi_relevan(pertanyaan)

            if konteks.strip() == "":
                st.warning("Belum ditemukan materi yang relevan. Pastikan admin sudah upload materi.")
                return

            prompt = f"""
            Kamu adalah AI Tutor TKA.

            Gunakan materi berikut sebagai sumber utama.

            Materi:
            {konteks}

            Pertanyaan siswa:
            {pertanyaan}

            Aturan:

            1. Jawab berdasarkan konsep yang terdapat pada materi.
            2. Jika siswa memberikan angka atau contoh baru yang tidak muncul persis di materi, tetapi konsepnya ada di materi, maka terapkan konsep tersebut untuk menjawab.
            3. Jangan mengarang konsep baru yang tidak terdapat pada materi.
            4. Jika konsep dasar tidak ditemukan dalam materi, katakan:
               "Materi yang diupload belum cukup untuk menjawab pertanyaan ini."
            5. Berikan langkah pengerjaan jika pertanyaan berupa soal.

            Jawaban:
            """

            jawaban = ask_llm(prompt)

            st.subheader("Jawaban AI Tutor")
            st.write(jawaban)


# =========================
# HALAMAN LEARNING PLANNER
# =========================
def page_learning_planner():
    st.title("🎯 Learning Planner AI")

    conn = connect_db()

    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        st.error("Data siswa tidak ditemukan.")
        conn.close()
        return

    mapel_dipilih = st.selectbox(
        "Pilih mapel untuk dibuatkan rencana belajar",
        MAPEL_OPTIONS
    )

    df = pd.read_sql_query("""
    SELECT mapel, nilai, benar, total, rekomendasi, created_at
    FROM hasil_tryout
    WHERE siswa_id = ? AND mapel = ?
    ORDER BY created_at DESC
    """, conn, params=(siswa_id, mapel_dipilih))

    conn.close()

    if df.empty:
        st.info(
            f"Belum ada hasil try out untuk mapel {mapel_dipilih}. Kerjakan Simulasi TKA mapel ini dulu agar AI bisa membuat rencana belajar.")
        return

    latest = df.iloc[0]

    nilai_rata_rata = round(df["nilai"].mean(), 2)
    nilai_tertinggi = df["nilai"].max()
    nilai_terendah = df["nilai"].min()
    jumlah_tryout = len(df)

    st.subheader("Ringkasan Hasil Try Out")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Mapel", mapel_dipilih)

    with col2:
        st.metric("Nilai Terakhir", latest["nilai"])

    with col3:
        st.metric("Rata-rata", nilai_rata_rata)

    with col4:
        st.metric("Jumlah Try Out", jumlah_tryout)

    st.write(f"Benar terakhir: **{latest['benar']} dari {latest['total']} soal**")
    st.write(f"Tanggal terakhir: **{latest['created_at']}**")

    st.divider()

    st.subheader("Riwayat Nilai Mapel Ini")
    st.dataframe(df, use_container_width=True)

    chart_df = df.sort_values("created_at")
    st.line_chart(chart_df.set_index("created_at")["nilai"])

    st.divider()

    target_nilai = st.number_input(
        "Target nilai yang ingin dicapai",
        min_value=0,
        max_value=100,
        value=85
    )

    waktu_belajar = st.selectbox(
        "Waktu belajar per hari",
        ["30 menit", "1 jam", "2 jam", "Lebih dari 2 jam"]
    )

    catatan_siswa = st.text_area(
        "Materi yang menurut kamu masih sulit pada mapel ini"
    )

    if st.button("Buat Rencana Belajar AI"):
        riwayat_ringkas = ""

        for _, row in df.head(5).iterrows():
            riwayat_ringkas += f"""
            - Tanggal: {row['created_at']}
              Nilai: {row['nilai']}
              Benar: {row['benar']} dari {row['total']}
              Rekomendasi sebelumnya: {row['rekomendasi']}
            """

        prompt = f"""
        Buatkan rencana belajar personal untuk siswa TKA.

        Mapel yang dipilih: {mapel_dipilih}

        Ringkasan performa siswa:
        - Jumlah try out pada mapel ini: {jumlah_tryout}
        - Nilai terakhir: {latest['nilai']}
        - Nilai rata-rata: {nilai_rata_rata}
        - Nilai tertinggi: {nilai_tertinggi}
        - Nilai terendah: {nilai_terendah}
        - Benar terakhir: {latest['benar']} dari {latest['total']} soal

        Riwayat try out terbaru:
        {riwayat_ringkas}

        Target nilai siswa: {target_nilai}
        Waktu belajar per hari: {waktu_belajar}
        Catatan materi sulit dari siswa: {catatan_siswa}

        Buat output:
        1. Ringkasan kondisi siswa pada mapel {mapel_dipilih}
        2. Gap antara nilai sekarang dan target
        3. Prioritas materi yang harus dipelajari
        4. Rencana belajar 7 hari
        5. Strategi latihan soal
        6. Saran penggunaan AI Tutor
        7. Motivasi singkat

        Gunakan bahasa Indonesia yang ramah, jelas, dan memotivasi.
        """

        with st.spinner("AI sedang membuat rencana belajar..."):
            hasil = ask_llm(prompt)

        st.subheader("🎯 Rencana Belajar AI")
        st.write(hasil)


# =========================
# PROGRESS SISWA
# =========================

def page_progress_saya():
    st.title("📈 Progress Saya")

    conn = connect_db()

    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        st.error("Data siswa tidak ditemukan.")
        conn.close()
        return

    tampilkan_penggunaan_simulasi_siswa(
        siswa_id,
        "📌 Penggunaan Simulasi TKA Kamu"
    )
    
    st.divider()

    df = pd.read_sql_query("""
    SELECT mapel, nilai, benar, total, rekomendasi, created_at
    FROM hasil_tryout
    WHERE siswa_id = ?
    ORDER BY created_at DESC
    """, conn, params=(siswa_id,))

    if df.empty:
        st.info("Belum ada data hasil try out.")
    else:
        st.dataframe(df, use_container_width=True)

        chart_df = df.sort_values("created_at")
        st.line_chart(chart_df.set_index("created_at")["nilai"])

        st.subheader("Rekomendasi Terbaru")
        st.write(df.iloc[0]["rekomendasi"])

    conn.close()


# =========================
# DASHBOARD ORANG TUA
# =========================

def page_orang_tua():
    role = st.session_state["role"]

    if role == "admin":
        st.title("📊 Hasil Try Out Siswa")
    else:
        st.title("👨‍👩‍👧 Dashboard Progress Anak")

    conn = connect_db()

    if role == "admin":
        tampilkan_rekap_penggunaan_simulasi_admin()
        st.divider()
        df_siswa = pd.read_sql_query("""
        SELECT id, nama, kelas, sekolah
        FROM siswa
        ORDER BY nama ASC
        """, conn)

        siswa_options = {"Semua Siswa": None}
        for _, row in df_siswa.iterrows():
            siswa_options[f"{row['id']} - {row['nama']} ({row['kelas']})"] = row["id"]

        filter_siswa_label = st.selectbox("Filter Siswa", list(siswa_options.keys()))
        filter_siswa_id = siswa_options[filter_siswa_label]

    elif role == "ortu":
        filter_siswa_id = get_siswa_id_by_ortu_username(st.session_state["username"])

        if filter_siswa_id is None:
            st.error("Data orang tua belum terhubung dengan siswa.")
            conn.close()
            return

        tampilkan_penggunaan_simulasi_siswa(
            filter_siswa_id,
            "📌 Penggunaan Simulasi TKA Anak"
        )
    
        st.divider()

    else:
        st.error("Akses tidak valid.")
        conn.close()
        return

    col1, col2, col3 = st.columns(3)

    with col1:
        filter_mapel = st.selectbox("Filter Mapel", ["Semua Mapel"] + MAPEL_OPTIONS)

    with col2:
        tanggal_mulai = st.date_input("Tanggal Mulai", value=None)

    with col3:
        tanggal_selesai = st.date_input("Tanggal Selesai", value=None)

    query = """
    SELECT 
        siswa.nama AS nama_siswa,
        siswa.kelas,
        siswa.sekolah,
        hasil_tryout.mapel,
        hasil_tryout.nilai,
        hasil_tryout.benar,
        hasil_tryout.total,
        hasil_tryout.rekomendasi,
        hasil_tryout.created_at
    FROM hasil_tryout
    JOIN siswa ON hasil_tryout.siswa_id = siswa.id
    WHERE 1=1
    """

    params = []

    if role == "ortu":
        query += " AND hasil_tryout.siswa_id = ?"
        params.append(filter_siswa_id)

    if role == "admin" and filter_siswa_id is not None:
        query += " AND hasil_tryout.siswa_id = ?"
        params.append(filter_siswa_id)

    if filter_mapel != "Semua Mapel":
        query += " AND hasil_tryout.mapel = ?"
        params.append(filter_mapel)

    if tanggal_mulai:
        query += " AND DATE(hasil_tryout.created_at) >= ?"
        params.append(str(tanggal_mulai))

    if tanggal_selesai:
        query += " AND DATE(hasil_tryout.created_at) <= ?"
        params.append(str(tanggal_selesai))

    query += " ORDER BY hasil_tryout.created_at DESC"

    df = pd.read_sql_query(query, conn, params=params)

    conn.close()

    if df.empty:
        st.info("Belum ada data hasil try out sesuai filter.")
        return

    st.dataframe(df, use_container_width=True)

    st.divider()

    col_metric1, col_metric2, col_metric3 = st.columns(3)

    with col_metric1:
        st.metric("Jumlah Try Out", len(df))

    with col_metric2:
        st.metric("Rata-rata Nilai", round(df["nilai"].mean(), 2))

    with col_metric3:
        st.metric("Nilai Tertinggi", df["nilai"].max())

    chart_df = df.sort_values("created_at")
    st.line_chart(chart_df.set_index("created_at")["nilai"])

    st.subheader("Rekomendasi Terbaru")
    st.write(df.iloc[0]["rekomendasi"])

    st.divider()
    st.subheader("Export Data")

    excel_file = export_excel(df)
    pdf_file = export_pdf(df, title="Laporan Hasil Try Out")

    col_exp1, col_exp2 = st.columns(2)

    with col_exp1:
        st.download_button(
            label="Download Excel",
            data=excel_file,
            file_name="hasil_tryout.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

    with col_exp2:
        st.download_button(
            label="Download PDF",
            data=pdf_file,
            file_name="hasil_tryout.pdf",
            mime="application/pdf",
            use_container_width=True
        )


# =========================
# HALAMAN GENERATE SOAL
# =========================
def page_generate_soal_materi():
    st.title("🤖 Generate Soal dari Materi")

    conn = connect_db()

    df_materi = pd.read_sql_query("""
    SELECT id, judul, filename, isi_teks
    FROM materi
    ORDER BY created_at DESC
    """, conn)

    if df_materi.empty:
        st.warning("Belum ada materi. Upload materi dulu di menu Upload Materi.")
        conn.close()
        return

    materi_options = {
        f"{row['id']} - {row['judul']}": row["id"]
        for _, row in df_materi.iterrows()
    }

    selected_label = st.selectbox("Pilih Materi", list(materi_options.keys()))
    selected_id = materi_options[selected_label]

    materi = df_materi[df_materi["id"] == selected_id].iloc[0]

    mapel = st.selectbox("Mapel", MAPEL_OPTIONS)
    topik = st.text_input("Topik Soal", value=materi["judul"])
    level = st.selectbox("Level Kesulitan", ["Mudah", "Sedang", "Sulit"])
    jumlah_soal = st.number_input("Jumlah soal", min_value=1, max_value=10, value=5)

    st.divider()

    st.subheader("Preview Materi")
    st.text_area("Isi materi", materi["isi_teks"][:3000], height=200)

    if st.button("Generate Soal dari Materi"):
        with st.spinner("AI sedang membuat soal. Mohon tunggu..."):
            try:
                materi_ringkas = materi["isi_teks"][:5000]

                prompt = f"""
                Buatkan {jumlah_soal} soal pilihan ganda TKA berdasarkan materi berikut.

                Materi:
                {materi_ringkas}

                Ketentuan:
                - Mapel: {mapel}
                - Topik: {topik}
                - Level kesulitan: {level}
                - Setiap soal memiliki 4 opsi: A, B, C, D
                - Jawaban benar harus A/B/C/D
                - Pembahasan singkat
                - Jangan keluar dari materi

                Kembalikan JSON valid saja tanpa markdown.

                Format:
                [
                  {{
                    "mapel": "{mapel}",
                    "topik": "{topik}",
                    "level": "{level}",
                    "pertanyaan": "teks soal",
                    "opsi_a": "pilihan A",
                    "opsi_b": "pilihan B",
                    "opsi_c": "pilihan C",
                    "opsi_d": "pilihan D",
                    "jawaban": "A",
                    "pembahasan": "pembahasan singkat"
                  }}
                ]
                """

                hasil = ask_llm(prompt)

                if not hasil or hasil.strip() == "":
                    st.error("AI tidak mengembalikan hasil.")
                    return

                st.session_state["generated_soal_json"] = hasil
                st.session_state["generated_soal_mapel"] = mapel
                st.session_state["generated_soal_topik"] = topik
                st.session_state["generated_soal_level"] = level

                st.subheader("Hasil Generate AI")
                st.text_area("Output AI", hasil, height=300)

            except Exception as e:
                st.error("Generate soal gagal.")
                st.code(str(e))

        st.session_state["generated_soal_json"] = hasil
        st.session_state["generated_soal_mapel"] = mapel
        st.session_state["generated_soal_topik"] = topik
        st.session_state["generated_soal_level"] = level

    if "generated_soal_json" in st.session_state:
        if st.button("Simpan Semua ke Bank Soal"):
            import json

            try:
                data_soal = json.loads(st.session_state["generated_soal_json"])

                cur = conn.cursor()

                sukses = 0

                for soal in data_soal:
                    cur.execute("""
                    INSERT INTO bank_soal
                    (mapel, topik, level, pertanyaan, opsi_a, opsi_b, opsi_c, opsi_d, jawaban, pembahasan)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        soal.get("mapel", st.session_state["generated_soal_mapel"]),
                        soal.get("topik", st.session_state["generated_soal_topik"]),
                        soal.get("level", st.session_state["generated_soal_level"]),
                        soal.get("pertanyaan", ""),
                        soal.get("opsi_a", ""),
                        soal.get("opsi_b", ""),
                        soal.get("opsi_c", ""),
                        soal.get("opsi_d", ""),
                        soal.get("jawaban", ""),
                        soal.get("pembahasan", "")
                    ))

                    sukses += 1

                conn.commit()
                st.success(f"{sukses} soal berhasil disimpan ke Bank Soal.")

            except Exception as e:
                st.error("Gagal menyimpan. Pastikan output AI berupa JSON valid.")
                st.code(str(e))

    conn.close()


# =====================
# EXPORT HASIL TRY OUT KE FILE
# ====================
def export_excel(df):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Hasil Try Out")

    output.seek(0)
    return output


def export_pdf(df, title="Laporan Hasil Try Out"):
    output = BytesIO()

    doc = SimpleDocTemplate(output, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 12))

    if df.empty:
        elements.append(Paragraph("Tidak ada data.", styles["Normal"]))
    else:
        data = [df.columns.tolist()] + df.astype(str).values.tolist()

        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        elements.append(table)

    doc.build(elements)
    output.seek(0)
    return output


# ===============================
# BERSIHKAN TEKS UNTUK PDF
# ===============================
def bersihkan_teks_pdf(teks):
    if pd.isna(teks):
        return "-"

    teks = str(teks)

    # Bersihkan markdown dari AI
    teks = teks.replace("###", "")
    teks = teks.replace("####", "")
    teks = teks.replace("##", "")
    teks = teks.replace("#", "")
    teks = teks.replace("**", "")
    teks = teks.replace("__", "")
    teks = teks.replace("`", "")

    # Ganti karakter yang sering bikin PDF aneh
    teks = teks.replace("–", "-")
    teks = teks.replace("—", "-")
    teks = teks.replace("•", "-")
    teks = teks.replace("■", "")
    teks = teks.replace("“", '"')
    teks = teks.replace("”", '"')
    teks = teks.replace("’", "'")
    teks = teks.replace("‘", "'")

    # Buang emoji / karakter yang tidak didukung font default PDF
    teks = teks.encode("latin-1", "ignore").decode("latin-1")

    # Rapikan spasi
    teks = re.sub(r"\n{3,}", "\n\n", teks)
    teks = re.sub(r"[ \t]+", " ", teks)

    return teks.strip()


def paragraph_pdf(teks, style):
    teks = bersihkan_teks_pdf(teks)
    teks = html.escape(teks)
    teks = teks.replace("\n", "<br/>")
    return Paragraph(teks, style)


def export_pdf(df, title="Laporan Hasil Try Out"):
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=1 * cm,
        leftMargin=1 * cm,
        topMargin=1 * cm,
        bottomMargin=1 * cm
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=14
    )

    header_style = ParagraphStyle(
        "TableHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        leading=10,
        alignment=TA_CENTER
    )

    cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10
    )

    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        spaceBefore=12,
        spaceAfter=8
    )

    detail_title_style = ParagraphStyle(
        "DetailTitle",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=13,
        spaceBefore=10,
        spaceAfter=4
    )

    detail_style = ParagraphStyle(
        "DetailText",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        spaceAfter=8
    )

    elements = []

    elements.append(Paragraph(title, title_style))

    if df.empty:
        elements.append(Paragraph("Tidak ada data hasil try out.", cell_style))
        doc.build(elements)
        output.seek(0)
        return output

    # =========================
    # TABEL RINGKAS
    # =========================
    elements.append(Paragraph("Ringkasan Hasil Try Out", section_style))

    data = [[
        Paragraph("No", header_style),
        Paragraph("Nama Siswa", header_style),
        Paragraph("Kelas", header_style),
        Paragraph("Sekolah", header_style),
        Paragraph("Mapel", header_style),
        Paragraph("Nilai", header_style),
        Paragraph("Benar/Total", header_style),
        Paragraph("Tanggal", header_style),
    ]]

    for i, row in df.iterrows():
        tanggal = row["created_at"] if "created_at" in df.columns else "-"
        benar_total = f"{row['benar']}/{row['total']}"

        data.append([
            Paragraph(str(len(data)), cell_style),
            Paragraph(bersihkan_teks_pdf(row["nama_siswa"]), cell_style),
            Paragraph(bersihkan_teks_pdf(row["kelas"]), cell_style),
            Paragraph(bersihkan_teks_pdf(row["sekolah"]), cell_style),
            Paragraph(bersihkan_teks_pdf(row["mapel"]), cell_style),
            Paragraph(str(row["nilai"]), cell_style),
            Paragraph(benar_total, cell_style),
            Paragraph(bersihkan_teks_pdf(tanggal), cell_style),
        ])

    table = Table(
        data,
        colWidths=[
            1.0 * cm,
            4.0 * cm,
            2.0 * cm,
            4.0 * cm,
            3.0 * cm,
            2.0 * cm,
            2.3 * cm,
            4.5 * cm,
        ],
        repeatRows=1
    )

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (5, 1), (6, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 14))

    # =========================
    # DETAIL REKOMENDASI
    # =========================

    elements.append(Paragraph("Detail Rekomendasi AI", section_style))

    for nomor, (_, row) in enumerate(df.iterrows(), start=1):
        nama = bersihkan_teks_pdf(row["nama_siswa"])
        mapel = bersihkan_teks_pdf(row["mapel"])
        nilai = row["nilai"]
        tanggal = bersihkan_teks_pdf(row["created_at"])

        judul_detail = f"{nomor}. {nama} - {mapel} - Nilai {nilai} - {tanggal}"

        elements.append(Paragraph(judul_detail, detail_title_style))

        if "rekomendasi" in df.columns:
            elements.append(paragraph_pdf(row["rekomendasi"], detail_style))
        else:
            elements.append(Paragraph("-", detail_style))

        elements.append(Spacer(1, 8))

    doc.build(elements)

    output.seek(0)
    return output


# ================================
# COBA GRATIS TRY OUT
# ================================
def page_tryout_gratis():
    st.title("🧪 Coba Try Out Gratis")

    JUMLAH_SOAL_GRATIS = 20

    st.info(
        "Fitur ini dapat dicoba tanpa login. "
        "Jumlah soal demo dibatasi 5 soal dan hasil tidak disimpan. "
        "Untuk mendapat try out adaptif penuh, jumlah soal fleksibel, progress tersimpan, "
        "learning planner, dan pantauan orang tua, silakan register terlebih dahulu."
    )

    if not st.session_state.get("free_trial_identitas_ok", False):
        st.subheader("Isi Data Sebelum Coba Gratis")
    
        with st.form("form_identitas_trial"):
            nama_trial = st.text_input("Nama")
            email_trial = st.text_input("Email")
            no_hp_trial = st.text_input("No HP / WhatsApp")
    
            submit_identitas = st.form_submit_button("Mulai Try Out Gratis")
    
            if submit_identitas:
                if not nama_trial or not email_trial or not no_hp_trial:
                    st.warning("Nama, email, dan no HP wajib diisi.")
                    return
    
                if sudah_pernah_tryout_gratis(email_trial, no_hp_trial):
                    st.warning(
                        "Data ini sudah pernah digunakan untuk mencoba try out gratis. "
                        "Untuk lanjut menggunakan fitur Smart TKA, silakan registrasi berbayar."
                    )
    
                    st.session_state["trial_nama"] = nama_trial
                    st.session_state["trial_email"] = email_trial
                    st.session_state["trial_no_hp"] = no_hp_trial
                    st.session_state["halaman_awal"] = "register_bayar"
                    st.rerun()
    
                st.session_state["trial_nama"] = nama_trial
                st.session_state["trial_email"] = email_trial
                st.session_state["trial_no_hp"] = no_hp_trial
                st.session_state["free_trial_identitas_ok"] = True
    
                st.rerun()
    
        return
    
    mapel_dipilih = st.selectbox(
        "Pilih Mapel",
        MAPEL_OPTIONS,
        key="free_mapel"
    )

    st.caption(
        f"Demo gratis untuk mapel {mapel_dipilih} berisi {JUMLAH_SOAL_GRATIS} soal."
    )

    conn = connect_db()

    # Buat paket otomatis saat pertama kali buka halaman
    # atau saat mapel diganti
    perlu_buat_paket = (
            "free_soal_ids" not in st.session_state
            or st.session_state.get("free_mapel_aktif") != mapel_dipilih
    )

    if perlu_buat_paket:
        df_paket = pd.read_sql_query("""
        SELECT id
        FROM bank_soal
        WHERE mapel = ?
        ORDER BY RANDOM()
        LIMIT ?
        """, conn, params=(mapel_dipilih, JUMLAH_SOAL_GRATIS))

        if df_paket.empty:
            st.warning(f"Belum ada soal untuk mapel {mapel_dipilih}.")
            conn.close()
            return

        st.session_state["free_soal_ids"] = df_paket["id"].tolist()
        st.session_state["free_mapel_aktif"] = mapel_dipilih

    soal_ids = st.session_state["free_soal_ids"]

    if not soal_ids:
        st.warning("Belum ada paket soal.")
        conn.close()
        return

    placeholders = ",".join(["?"] * len(soal_ids))

    df = pd.read_sql_query(f"""
    SELECT *
    FROM bank_soal
    WHERE id IN ({placeholders})
    """, conn, params=soal_ids)

    conn.close()

    if df.empty:
        st.warning("Soal tidak ditemukan.")
        return

    # Supaya urutan soal tetap sama seperti paket yang dibuat
    df["urutan_paket"] = df["id"].apply(lambda x: soal_ids.index(x))
    df = df.sort_values("urutan_paket")

    st.info(
        f"Mapel: {mapel_dipilih} | Jumlah soal demo: {len(df)}"
    )

    jawaban_user = {}

    with st.form("form_tryout_gratis"):
        for nomor, (_, row) in enumerate(df.iterrows(), start=1):
            st.markdown(f"### Soal {nomor}")
            st.write(row["pertanyaan"])

            topik = row["topik"] if "topik" in row and pd.notna(row["topik"]) else "-"
            level = row["level"] if "level" in row and pd.notna(row["level"]) else "-"

            st.caption(f"Topik: {topik} | Level: {level}")

            pilihan = {
                "A": row["opsi_a"],
                "B": row["opsi_b"],
                "C": row["opsi_c"],
                "D": row["opsi_d"],
            }

            jawaban_user[row["id"]] = st.radio(
                "Pilih jawaban",
                options=["A", "B", "C", "D"],
                index=None,
                format_func=lambda x: f"{x}. {pilihan[x]}",
                key=f"free_soal_{row['id']}"
            )

        submit = st.form_submit_button("Submit Try Out Gratis")

    if submit:
        benar = 0
        total = len(df)

        topik_benar = []
        topik_salah = []
        pembahasan_ringkas = []

        for _, row in df.iterrows():
            jawaban_siswa = jawaban_user[row["id"]]
            jawaban_benar = row["jawaban"]

            topik = row["topik"] if "topik" in row and pd.notna(row["topik"]) else "-"

            if jawaban_siswa == jawaban_benar:
                benar += 1
                topik_benar.append(topik)
                status = "Benar"
            else:
                topik_salah.append(topik)
                status = "Salah"

            pembahasan_ringkas.append({
                "Pertanyaan": row["pertanyaan"],
                "Jawaban Kamu": jawaban_siswa,
                "Jawaban Benar": jawaban_benar,
                "Status": status
            })

        nilai = round((benar / total) * 100, 2)

        st.success(f"Nilai demo kamu: {nilai}")
        st.write(f"Benar: **{benar} dari {total} soal**")

        if not st.session_state.get("free_trial_sudah_dicatat", False):
            catat_tryout_gratis(
                st.session_state.get("trial_nama", ""),
                st.session_state.get("trial_email", ""),
                st.session_state.get("trial_no_hp", ""),
                mapel_dipilih,
                nilai,
                benar,
                total
            )

            st.session_state["free_trial_sudah_dicatat"] = True

        st.subheader("📌 Ringkasan Jawaban")
        st.dataframe(pd.DataFrame(pembahasan_ringkas), use_container_width=True)

        st.subheader("✅ Topik yang Sudah Dikuasai")
        st.write(list(set(topik_benar)) if topik_benar else "-")

        st.subheader("⚠️ Topik yang Perlu Dipelajari")
        st.write(list(set(topik_salah)) if topik_salah else "-")

        prompt = f"""
        Buat diagnosis singkat untuk calon pengguna yang mencoba try out gratis.

        Mapel: {mapel_dipilih}
        Nilai: {nilai}
        Benar: {benar} dari {total}

        Topik yang dikuasai:
        {list(set(topik_benar))}

        Topik yang masih salah:
        {list(set(topik_salah))}

        Buat output:
        1. Ringkasan kemampuan
        2. Materi yang perlu dipelajari
        3. Saran belajar singkat
        4. Ajakan untuk register agar progress belajar bisa tersimpan

        Gunakan bahasa Indonesia yang ramah dan memotivasi.
        """

        with st.spinner("AI sedang membuat diagnosis singkat..."):
            diagnosis = ask_llm(prompt)

        st.subheader("🤖 Diagnosis Singkat AI")
        st.write(diagnosis)

        st.divider()

        st.info(
            "Hasil demo ini belum tersimpan. "
            "Untuk menyimpan progress, mendapatkan soal adaptif, learning planner, "
            "adaptive learning, dan pantauan orang tua, silakan register."
        )

        st.button(
            "Lanjut Registrasi Berbayar Rp100.000",
            on_click=redirect_ke_register_bayar,
            use_container_width=True
        )


# =========================
# MENGAMBIL TOPIK PRIORITAS UNTUK JADWAL BELAJAR
# =========================
def ambil_topik_prioritas(siswa_id, mapel):
    conn = connect_db()

    df = pd.read_sql_query("""
    SELECT 
        topik,
        level,
        COUNT(*) AS total_dikerjakan,
        SUM(CASE WHEN is_benar = 1 THEN 1 ELSE 0 END) AS total_benar,
        SUM(CASE WHEN is_benar = 0 THEN 1 ELSE 0 END) AS total_salah,
        ROUND(AVG(is_benar) * 100, 2) AS akurasi
    FROM hasil_tryout_detail
    WHERE siswa_id = ? AND mapel = ?
    GROUP BY topik, level
    ORDER BY akurasi ASC, total_salah DESC, total_dikerjakan DESC
    """, conn, params=(siswa_id, mapel))

    conn.close()

    return df


# =========================
# FITUR HALAMAN JADWAL BELAJAR OTOMATIS
# =========================
def page_jadwal_belajar_otomatis():
    st.title("📅 Jadwal Belajar Otomatis")

    conn = connect_db()

    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        st.error("Data siswa tidak ditemukan.")
        conn.close()
        return

    st.info(
        "Sistem akan membuat jadwal belajar otomatis berdasarkan hasil try out, "
        "khususnya dari topik yang masih lemah."
    )

    mapel_dipilih = st.selectbox(
        "Pilih Mapel",
        MAPEL_OPTIONS
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        jumlah_hari = st.selectbox(
            "Durasi Jadwal",
            [7, 14, 21, 30],
            index=0
        )

    with col2:
        target_nilai = st.number_input(
            "Target Nilai",
            min_value=0,
            max_value=100,
            value=85
        )

    with col3:
        waktu_per_hari = st.selectbox(
            "Waktu Belajar per Hari",
            ["30 menit", "1 jam", "2 jam"]
        )

    col4, col5, col6 = st.columns(3)
    with col4:
        tanggal_mulai = st.date_input("Tanggal Mulai")

    with col5:
        jam_belajar = st.time_input("Mulai Jam Belajar Harian")

    with col6:
        reminder_aktif = st.selectbox(
            "Aktifkan Reminder?",
            ["Ya", "Tidak"]
        )

    df_prioritas = ambil_topik_prioritas(siswa_id, mapel_dipilih)

    if df_prioritas.empty:
        st.warning(
            f"Belum ada data detail try out untuk mapel {mapel_dipilih}. "
            "Kerjakan try out dulu agar sistem bisa membuat jadwal otomatis."
        )
        conn.close()
        return

    st.subheader("📊 Topik Prioritas dari Hasil Try Out")
    st.dataframe(df_prioritas, use_container_width=True)

    if st.button("Buat Jadwal Belajar Otomatis"):
        df_soal_salah = ambil_soal_salah_siswa(siswa_id, mapel_dipilih)

        daftar_topik_ai = analisis_kelemahan_spesifik_ai(
            df_soal_salah,
            mapel_dipilih
        )

        topik_prioritas_text = ""

        if daftar_topik_ai:
            daftar_topik = daftar_topik_ai

            for item in daftar_topik:
                topik_prioritas_text += f"""
                - Topik spesifik: {item.get("topik_spesifik", "-")}
                  Konteks kelemahan: {item.get("konteks_kelemahan", "-")}
                  Materi harian: {item.get("materi_harian", "-")}
                  Latihan disarankan: {item.get("latihan_disarankan", "-")}
                  Level: {item.get("level", "Mudah")}
                """

        else:
            topik_lemah = df_prioritas[df_prioritas["akurasi"] < 80]

            if topik_lemah.empty:
                topik_lemah = df_prioritas.head(3)

            daftar_topik = []

            for _, row in topik_lemah.iterrows():
                daftar_topik.append({
                    "topik_spesifik": row["topik"],
                    "konteks_kelemahan": f"Akurasi pada topik ini masih {row['akurasi']}%.",
                    "materi_harian": f"Pelajari ulang materi {row['topik']}.",
                    "latihan_disarankan": f"Kerjakan latihan soal level {row['level']} tentang {row['topik']}.",
                    "level": row["level"]
                })

            topik_prioritas_text = topik_lemah.to_string(index=False)

        aktivitas_template = [
            "Review konsep dasar",
            "Membaca materi dan mencatat poin penting",
            "Latihan soal mudah",
            "Latihan soal sedang",
            "Bahas kesalahan dari try out sebelumnya",
            "Latihan soal campuran",
            "Evaluasi mandiri dan rangkum materi"
        ]

        cur = conn.cursor()

        try:
            prompt = f"""
            Buat catatan singkat untuk jadwal belajar siswa.

            Mapel: {mapel_dipilih}
            Target nilai: {target_nilai}
            Durasi jadwal: {jumlah_hari} hari
            Waktu belajar per hari: {waktu_per_hari}

            Topik prioritas:
            {topik_prioritas_text}

            Buat catatan motivasi dan strategi belajar singkat dalam bahasa Indonesia.
            """

            with st.spinner("AI sedang membuat catatan strategi belajar..."):
                catatan_ai = ask_llm(prompt)

            cur.execute("""
            INSERT INTO jadwal_belajar (
                siswa_id, mapel, tanggal_mulai, jumlah_hari, target_nilai,
                waktu_per_hari, jam_belajar, reminder_aktif, catatan_ai
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                siswa_id,
                mapel_dipilih,
                str(tanggal_mulai),
                jumlah_hari,
                target_nilai,
                waktu_per_hari,
                str(jam_belajar),
                reminder_aktif,
                catatan_ai
            ))

            jadwal_id = cur.lastrowid

            for i in range(jumlah_hari):
                hari_ke = i + 1
                tanggal = tanggal_mulai + pd.Timedelta(days=i)

                topik_data = daftar_topik[i % len(daftar_topik)]

                topik_spesifik = topik_data.get("topik_spesifik", "-")
                konteks_kelemahan = topik_data.get("konteks_kelemahan", "-")
                materi_harian = topik_data.get("materi_harian", "-")
                latihan_disarankan = topik_data.get("latihan_disarankan", "-")
                level = topik_data.get("level", "Mudah")

                fokus_belajar = f"Fokus pada {topik_spesifik}"

                aktivitas_final = (
                    f"Pelajari {topik_spesifik}. "
                    f"Konteks kelemahan: {konteks_kelemahan}"
                )

                cur.execute("""
                INSERT INTO jadwal_belajar_detail (
                    jadwal_id, hari_ke, tanggal, topik, level,
                    topik_spesifik, konteks_kelemahan,
                    fokus_belajar, materi_harian, latihan_disarankan,
                    aktivitas, status
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'belum')
                """, (
                    jadwal_id,
                    hari_ke,
                    str(tanggal),
                    topik_spesifik,
                    level,
                    topik_spesifik,
                    konteks_kelemahan,
                    fokus_belajar,
                    materi_harian,
                    latihan_disarankan,
                    aktivitas_final
                ))

            conn.commit()

            st.success("Jadwal belajar otomatis berhasil dibuat.")
            st.rerun()

        except Exception as e:
            conn.rollback()
            st.error("Gagal membuat jadwal belajar.")
            st.code(str(e))

    st.divider()

    st.subheader("📚 Jadwal Belajar Saya")

    df_jadwal = pd.read_sql_query("""
    SELECT *
    FROM jadwal_belajar
    WHERE siswa_id = ?
    ORDER BY created_at DESC
    """, conn, params=(siswa_id,))

    if df_jadwal.empty:
        st.info("Belum ada jadwal belajar.")
        conn.close()
        return

    pilihan_jadwal = st.selectbox(
        "Pilih Jadwal",
        df_jadwal["id"].tolist(),
        format_func=lambda x: f"Jadwal #{x}"
    )

    data_jadwal = df_jadwal[df_jadwal["id"] == pilihan_jadwal].iloc[0]

    st.write(f"**Mapel:** {data_jadwal['mapel']}")
    st.write(f"**Tanggal Mulai:** {data_jadwal['tanggal_mulai']}")
    st.write(f"**Durasi:** {data_jadwal['jumlah_hari']} hari")
    st.write(f"**Target Nilai:** {data_jadwal['target_nilai']}")
    st.write(f"**Waktu Belajar:** {data_jadwal['waktu_per_hari']}")

    st.subheader("🤖 Catatan Strategi AI")
    st.write(data_jadwal["catatan_ai"])

    df_detail = pd.read_sql_query("""
    SELECT *
    FROM jadwal_belajar_detail
    WHERE jadwal_id = ?
    ORDER BY hari_ke ASC
    """, conn, params=(pilihan_jadwal,))

    st.subheader("📅 Detail Jadwal Harian")

    kolom_tampil = [
        "hari_ke",
        "tanggal",
        "topik_spesifik",
        "level",
        "konteks_kelemahan",
        "materi_harian",
        "latihan_disarankan",
        "status"
    ]

    kolom_tersedia = [kol for kol in kolom_tampil if kol in df_detail.columns]

    st.dataframe(df_detail[kolom_tersedia], use_container_width=True)

    st.divider()

    st.subheader("📖 Panduan Belajar Hari per Hari")

    for _, row in df_detail.iterrows():
        topik_tampil = row["topik_spesifik"] if "topik_spesifik" in row and pd.notna(row["topik_spesifik"]) else row[
            "topik"]

        with st.expander(f"Hari ke-{row['hari_ke']} | {row['tanggal']} | {topik_tampil}"):
            st.write(f"**Topik Spesifik:** {topik_tampil}")
            st.write(f"**Level:** {row['level']}")

            if "konteks_kelemahan" in row:
                st.write(f"**Konteks Kelemahan:** {row['konteks_kelemahan']}")

            if "fokus_belajar" in row:
                st.write(f"**Fokus Belajar:** {row['fokus_belajar']}")

            if "materi_harian" in row:
                st.write(f"**Materi Hari Ini:** {row['materi_harian']}")

            if "latihan_disarankan" in row:
                st.write(f"**Latihan Disarankan:** {row['latihan_disarankan']}")

            st.write(f"**Status:** {row['status']}")

    st.divider()

    st.subheader("✅ Update Status Belajar")

    df_belum = df_detail[df_detail["status"] != "selesai"]

    if df_belum.empty:
        st.success("Semua jadwal belajar sudah selesai. Mantap!")
    else:
        pilihan_detail = st.selectbox(
            "Pilih Hari yang Sudah Dikerjakan",
            df_belum["id"].tolist(),
            format_func=lambda x: f"Hari ke-{int(df_belum[df_belum['id'] == x].iloc[0]['hari_ke'])}"
        )

        if st.button("Tandai Selesai"):
            cur = conn.cursor()

            cur.execute("""
            UPDATE jadwal_belajar_detail
            SET status = 'selesai'
            WHERE id = ?
            """, (pilihan_detail,))

            conn.commit()
            st.success("Status belajar berhasil diperbarui.")
            st.rerun()

    conn.close()


# ===========================
# AMBIL SOAL YANG SALAH
# ===========================
def ambil_soal_salah_siswa(siswa_id, mapel):
    conn = connect_db()

    df = pd.read_sql_query("""
    SELECT 
        d.topik,
        d.level,
        d.jawaban_siswa,
        d.jawaban_benar,
        b.pertanyaan,
        b.opsi_a,
        b.opsi_b,
        b.opsi_c,
        b.opsi_d
    FROM hasil_tryout_detail d
    JOIN bank_soal b ON d.soal_id = b.id
    WHERE d.siswa_id = ?
      AND d.mapel = ?
      AND d.is_benar = 0
    ORDER BY d.created_at DESC
    LIMIT 15
    """, conn, params=(siswa_id, mapel))

    conn.close()

    return df


# ===============================
# ANALISIS KELEMAHAN SPESIFIK
# ===============================
def analisis_kelemahan_spesifik_ai(df_salah, mapel):
    if df_salah.empty:
        return []

    daftar_soal = ""

    for i, row in df_salah.iterrows():
        daftar_soal += f"""
                        Soal:
                        {row['pertanyaan']}

                        Opsi:
                        A. {row['opsi_a']}
                        B. {row['opsi_b']}
                        C. {row['opsi_c']}
                        D. {row['opsi_d']}

                        Jawaban siswa: {row['jawaban_siswa']}
                        Jawaban benar: {row['jawaban_benar']}
                        Topik lama dari bank soal: {row['topik']}
                        Level: {row['level']}
                        ---
                        """
    prompt = f"""
            Kamu adalah sistem analisis adaptive learning.
            Tugasmu adalah membaca daftar soal yang dijawab salah oleh siswa, lalu menyimpulkan kelemahan belajar yang lebih spesifik.
            Mapel: {mapel}

            Data soal yang salah:
            {daftar_soal}
            Buat output dalam JSON valid saja, tanpa markdown, tanpa penjelasan tambahan.
            Format:
            [
              {{
                "topik_spesifik": "nama subtopik spesifik",
                "konteks_kelemahan": "penjelasan singkat kelemahan siswa",
                "materi_harian": "materi yang perlu dipelajari siswa",
                "latihan_disarankan": "jenis latihan yang disarankan",
                "level": "Mudah/Sedang/Sulit"
              }}
            ]

            Aturan:
            - Jangan pakai topik umum seperti "Matematika Kelas III".
            - Buat topik spesifik berdasarkan isi soal.
            - Contoh topik spesifik:
              - Membaca Bilangan Cacah Tiga Angka
              - Nilai Tempat Ratusan Puluhan Satuan
              - Membandingkan Bilangan Cacah
              - Perkalian Dasar
              - Sinonim
              - Antonim
              - Main Idea
              - Vocabulary
            - Gabungkan soal yang kelemahannya mirip menjadi satu topik spesifik.
            - Maksimal 5 topik spesifik.
            """

    hasil = ask_llm(prompt)

    try:
        hasil_bersih = hasil.strip()
        hasil_bersih = hasil_bersih.replace("```json", "")
        hasil_bersih = hasil_bersih.replace("```", "")
        hasil_bersih = hasil_bersih.strip()

        data = json.loads(hasil_bersih)
        return data

    except Exception:
        return []


# =======================
# DETAIL BELAJAR HARIAN
# =======================
def buat_detail_belajar_harian(topik, level, akurasi, hari_ke):
    pola_hari = (hari_ke - 1) % 7

    if akurasi < 50:
        kondisi = "dasar"
    elif akurasi < 80:
        kondisi = "penguatan"
    else:
        kondisi = "lanjutan"

    if pola_hari == 0:
        fokus_belajar = f"Memahami ulang konsep dasar {topik}"
        materi_harian = f"Pelajari kembali pengertian, contoh, dan langkah dasar pada topik {topik}."
        latihan_disarankan = f"Kerjakan 5-10 soal level {level} tentang {topik}."

    elif pola_hari == 1:
        fokus_belajar = f"Contoh soal dan pembahasan {topik}"
        materi_harian = f"Baca contoh soal {topik}, lalu pahami cara menentukan jawabannya."
        latihan_disarankan = f"Kerjakan soal latihan bertahap dari yang paling mudah tentang {topik}."

    elif pola_hari == 2:
        fokus_belajar = f"Latihan soal terarah {topik}"
        materi_harian = f"Fokus pada bagian {topik} yang masih sering salah saat try out."
        latihan_disarankan = f"Kerjakan minimal 10 soal {topik}, lalu catat soal yang salah."

    elif pola_hari == 3:
        fokus_belajar = f"Membahas kesalahan pada topik {topik}"
        materi_harian = f"Pelajari ulang soal-soal {topik} yang sebelumnya dijawab salah."
        latihan_disarankan = f"Tulis alasan kenapa jawaban sebelumnya salah, lalu coba ulangi soal sejenis."

    elif pola_hari == 4:
        fokus_belajar = f"Penguatan konsep {topik}"
        materi_harian = f"Buat rangkuman singkat tentang rumus, pola, atau konsep penting dalam {topik}."
        latihan_disarankan = f"Kerjakan latihan campuran level {level} tentang {topik}."

    elif pola_hari == 5:
        fokus_belajar = f"Simulasi mini topik {topik}"
        materi_harian = f"Kerjakan simulasi kecil khusus topik {topik} tanpa melihat catatan."
        latihan_disarankan = f"Kerjakan 10-15 soal {topik} dan hitung persentase benar."

    else:
        fokus_belajar = f"Evaluasi dan rangkuman {topik}"
        materi_harian = f"Tinjau ulang hasil belajar minggu ini untuk topik {topik}."
        latihan_disarankan = f"Buat catatan kesalahan dan tanyakan ke AI Tutor bagian yang masih belum paham."

    if kondisi == "dasar":
        aktivitas = f"Belajar dari dasar karena akurasi try out masih rendah. Fokus pada {topik} level {level}."
    elif kondisi == "penguatan":
        aktivitas = f"Perkuat pemahaman karena akurasi mulai berkembang. Fokus pada {topik} level {level}."
    else:
        aktivitas = f"Pertahankan kemampuan dan coba soal lebih menantang pada topik {topik} level {level}."

    return fokus_belajar, materi_harian, latihan_disarankan, aktivitas


# ==========================
# AMBIL REMINDER BELAJAR
# =========================
def ambil_reminder_belajar_siswa(siswa_id):
    conn = connect_db()

    hari_ini = str(date.today())

    df_hari_ini = pd.read_sql_query("""
    SELECT 
        jd.id,
        j.mapel,
        j.jam_belajar,
        jd.hari_ke,
        jd.tanggal,
        jd.topik,
        jd.topik_spesifik,
        jd.level,
        jd.fokus_belajar,
        jd.materi_harian,
        jd.latihan_disarankan,
        jd.status
    FROM jadwal_belajar_detail jd
    JOIN jadwal_belajar j ON jd.jadwal_id = j.id
    WHERE j.siswa_id = ?
      AND j.reminder_aktif = 'Ya'
      AND jd.tanggal = ?
      AND jd.status != 'selesai'
    ORDER BY jd.tanggal ASC, jd.hari_ke ASC
    """, conn, params=(siswa_id, hari_ini))

    df_terlewat = pd.read_sql_query("""
    SELECT 
        jd.id,
        j.mapel,
        j.jam_belajar,
        jd.hari_ke,
        jd.tanggal,
        jd.topik,
        jd.topik_spesifik,
        jd.level,
        jd.fokus_belajar,
        jd.materi_harian,
        jd.latihan_disarankan,
        jd.status
    FROM jadwal_belajar_detail jd
    JOIN jadwal_belajar j ON jd.jadwal_id = j.id
    WHERE j.siswa_id = ?
      AND j.reminder_aktif = 'Ya'
      AND jd.tanggal < ?
      AND jd.status != 'selesai'
    ORDER BY jd.tanggal ASC, jd.hari_ke ASC
    """, conn, params=(siswa_id, hari_ini))

    conn.close()

    return df_hari_ini, df_terlewat


# ===========================
# NOTIFIKASI REMINDER BELAJAR
# ===========================
def tampilkan_notifikasi_belajar():
    siswa_id = get_siswa_id_by_username(st.session_state["username"])

    if siswa_id is None:
        return

    df_hari_ini, df_terlewat = ambil_reminder_belajar_siswa(siswa_id)

    if not df_terlewat.empty:
        st.error(f"⚠️ Ada {len(df_terlewat)} jadwal belajar yang terlewat dan belum ditandai selesai.")

        with st.expander("Lihat jadwal yang terlewat"):
            for _, row in df_terlewat.iterrows():
                topik = row["topik_spesifik"] if pd.notna(row["topik_spesifik"]) else row["topik"]

                st.write(f"**Tanggal:** {row['tanggal']}")
                st.write(f"**Mapel:** {row['mapel']}")
                st.write(f"**Topik:** {topik}")
                st.write(f"**Fokus:** {row['fokus_belajar']}")
                st.write("---")

    if not df_hari_ini.empty:
        st.warning(f"📌 Kamu punya {len(df_hari_ini)} jadwal belajar hari ini.")

        with st.expander("Lihat jadwal belajar hari ini", expanded=True):
            for _, row in df_hari_ini.iterrows():
                topik = row["topik_spesifik"] if pd.notna(row["topik_spesifik"]) else row["topik"]

                st.write(f"**Jam belajar:** {row['jam_belajar']}")
                st.write(f"**Mapel:** {row['mapel']}")
                st.write(f"**Topik:** {topik}")
                st.write(f"**Level:** {row['level']}")
                st.write(f"**Fokus belajar:** {row['fokus_belajar']}")
                st.write(f"**Materi hari ini:** {row['materi_harian']}")
                st.write(f"**Latihan:** {row['latihan_disarankan']}")
                st.write("---")


# ========================
# CATAT VISITOR HALAMAN WEB
# =====================
def catat_pengunjung(halaman="Halaman Awal"):
    if "visitor_session_id" not in st.session_state:
        st.session_state["visitor_session_id"] = str(uuid.uuid4())

    session_id = st.session_state["visitor_session_id"]
    tanggal = str(date.today())

    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT COUNT(*)
    FROM visitor_log
    WHERE session_id = ?
    """, (session_id,))

    sudah_ada = cur.fetchone()[0]

    if sudah_ada == 0:
        cur.execute("""
        INSERT INTO visitor_log 
        (session_id, halaman, role, username, tanggal, is_login)
        VALUES (?, ?, ?, ?, ?, 0)
        """, (
            session_id,
            halaman,
            "guest",
            None,
            tanggal
        ))

        conn.commit()

    conn.close()

# =========================
# AMBIL STATISTIK PENGUNJUNG HALAMAN
# =========================
def ambil_statistik_pengunjung():
    conn = connect_db()

    hari_ini = str(date.today())

    total_pengunjung = pd.read_sql_query("""
    SELECT COUNT(DISTINCT session_id) AS total
    FROM visitor_log
    """, conn).iloc[0]["total"]

    pengunjung_hari_ini = pd.read_sql_query("""
    SELECT COUNT(DISTINCT session_id) AS total
    FROM visitor_log
    WHERE tanggal = ?
    """, conn, params=(hari_ini,)).iloc[0]["total"]

    login_hari_ini = pd.read_sql_query("""
    SELECT COUNT(DISTINCT session_id) AS total
    FROM visitor_log
    WHERE tanggal = ? AND is_login = 1
    """, conn, params=(hari_ini,)).iloc[0]["total"]

    guest_hari_ini = pd.read_sql_query("""
    SELECT COUNT(DISTINCT session_id) AS total
    FROM visitor_log
    WHERE tanggal = ? AND is_login = 0
    """, conn, params=(hari_ini,)).iloc[0]["total"]

    total_login = pd.read_sql_query("""
    SELECT COUNT(DISTINCT session_id) AS total
    FROM visitor_log
    WHERE is_login = 1
    """, conn).iloc[0]["total"]

    conn.close()

    return {
        "total_pengunjung": total_pengunjung,
        "pengunjung_hari_ini": pengunjung_hari_ini,
        "login_hari_ini": login_hari_ini,
        "guest_hari_ini": guest_hari_ini,
        "total_login": total_login
    }

# =========================
# CATAT YANG LOGIN
# =========================
def catat_login_pengunjung(username, role):
    if "visitor_session_id" not in st.session_state:
        st.session_state["visitor_session_id"] = str(uuid.uuid4())

    session_id = st.session_state["visitor_session_id"]
    tanggal = str(date.today())

    conn = connect_db()
    cur = conn.cursor()

    cur.execute("""
    SELECT COUNT(*)
    FROM visitor_log
    WHERE session_id = ?
    """, (session_id,))

    sudah_ada = cur.fetchone()[0]

    if sudah_ada == 0:
        cur.execute("""
        INSERT INTO visitor_log 
        (session_id, halaman, role, username, tanggal, is_login, login_at)
        VALUES (?, ?, ?, ?, ?, 1, CURRENT_TIMESTAMP)
        """, (
            session_id,
            "Login",
            role,
            username,
            tanggal
        ))
    else:
        cur.execute("""
        UPDATE visitor_log
        SET is_login = 1,
            username = ?,
            role = ?,
            login_at = CURRENT_TIMESTAMP
        WHERE session_id = ?
        """, (
            username,
            role,
            session_id
        ))

    conn.commit()
    conn.close()

# =========================
# TAMPILAN STATS COUNTER
# =========================
def tampilkan_stats_counter():
    stats = ambil_statistik_pengunjung()

    with st.container(border=True):
        st.subheader("📊 Statistik Pengunjung")

        st.write(f"**Total Pengunjung:** {int(stats['total_pengunjung'])}")
        st.write(f"**Pengunjung Hari Ini:** {int(stats['pengunjung_hari_ini'])}")
        st.write(f"**Login Hari Ini:** {int(stats['login_hari_ini'])}")
        st.write(f"**Guest Hari Ini:** {int(stats['guest_hari_ini'])}")
        st.write(f"**Total Pernah Login:** {int(stats['total_login'])}")

        st.caption("Guest = pengunjung yang membuka aplikasi tapi belum login.")


# =========================
# MAIN
# =========================

#init_db()
jalankan_init_db_sekali(DB_SCHEMA_VERSION)

if "username" not in st.session_state:
    catat_pengunjung("Halaman Awal")

    if st.session_state.get("halaman_awal") == "register_bayar":
        if st.button("← Kembali ke Halaman Awal"):
            st.session_state["halaman_awal"] = "awal"
            st.rerun()

        page_register_bayar()

    else:
        tab1, tab2, tab3, tab4 = st.tabs([
            "🔐 Login",
            "🧪 Coba Try Out Gratis",
            "💳 Register Berbayar",
            "🔎 Cek Status Registrasi"
        ])
        
        with tab1:
            col_login, col_stats = st.columns([2, 1])
        
            with col_login:
                page_login()
        
            with col_stats:
                st.markdown("<div style='height: 32px;'></div>", unsafe_allow_html=True)
                tampilkan_stats_counter()
        
        with tab2:
            page_tryout_gratis()
        
        with tab3:
            page_register_bayar()
        
        with tab4:
            page_cek_status_registrasi()
else:
    st.sidebar.title("📘 TKA Digital")
    st.sidebar.write(f"Login sebagai: **{st.session_state['nama']}**")

    role = st.session_state["role"]

    menu_options = ["Dashboard"]

    if role == "admin":
        menu_options += [
            "Bank Soal",
            "Upload Materi",
            "Generate Soal AI",
            "AI Tutor",
            "Hasil Try Out",
            "Data Siswa",
            "Data Orang Tua",
            "Verifikasi Register"
        ]

    elif role == "siswa":
        menu_options += [
            "Simulasi TKA",
            "AI Tutor",
            "Learning Planner",
            "Adaptive Learning",
            "Jadwal Belajar",
            "Progress Saya"
        ]

    elif role == "ortu":
        menu_options += [
            "Dashboard Orang Tua"
        ]

    if "menu" not in st.session_state:
        st.session_state["menu"] = "Dashboard"

    if st.session_state["menu"] not in menu_options:
        st.session_state["menu"] = "Dashboard"

    if "sidebar_menu" not in st.session_state:
        st.session_state["sidebar_menu"] = st.session_state.get("menu", "Dashboard")

    if st.session_state["sidebar_menu"] not in menu_options:
        st.session_state["sidebar_menu"] = "Dashboard"
    
    menu = st.sidebar.radio(
        "Menu",
        menu_options,
        index=menu_options.index(st.session_state["sidebar_menu"]),
        key="sidebar_menu"
    )
    
    st.session_state["menu"] = menu

    if st.sidebar.button("Logout"):
        logout()

    if menu == "Dashboard":
        page_dashboard()

    elif menu == "Bank Soal":
        page_bank_soal()

    elif menu == "Upload Materi":
        page_materi()

    elif menu == "Simulasi TKA":
        page_tryout()

    elif menu == "AI Tutor":
        page_ai_tutor()

    elif menu == "Generate Soal AI":
        page_generate_soal_materi()

    elif menu == "Learning Planner":
        page_learning_planner()

    elif menu == "Progress Saya":
        page_progress_saya()

    elif menu == "Dashboard Orang Tua":
        page_orang_tua()

    elif menu == "Hasil Try Out":
        page_orang_tua()

    elif menu == "Data Siswa":
        page_crud_siswa()

    elif menu == "Data Orang Tua":
        page_crud_ortu()

    elif menu == "Adaptive Learning":
        page_adaptive_learning()

    elif menu == "Verifikasi Register":
        page_verifikasi_register()

    elif menu == "Jadwal Belajar":
        page_jadwal_belajar_otomatis()
